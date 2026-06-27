"""
Build a scoped FSD briefing from canonical Informatica XML JSON.

This script is designed for JSON produced by `informatica_xml_to_json.py`, i.e.
node-based shape with keys: `tag`, `attributes`, `children`.

ENHANCED VERSION v2 — Zero-Loss Mapping Strategy Applied
---------------------------------------------------------
Implements all recommendations from zero_loss_mapping_strategy.md:

  1. Multi-phase processing with intermediate graph representation
  2. Persisted intermediate artifacts (temp files / durable cache)
  3. Incremental checkpointing with count/consistency validation
  4. Graph memoization keyed by (instance, port)
  5. Robust traversal — never prematurely terminates for opaque types
  6. Path-aware visited tracking — separates cycle detection from node caching
  7. Completeness validation — flags unresolved nodes, verifies all target fields
  8. Diagnostics — warnings for unresolved expressions, ports, missing connectors

Transformation-specific handling (carried from v1):
- Router, Joiner, Union with dot-suffix and multiple groups
- Java/SQL Transformations (opaque types with name-match fallback)
- Normalizer with generated ports
- Lookup (connected and unconnected)
- Mapplets with sub-mapping recursion
- Generated ports (RANKINDEX, DD_*, TCONTROL, etc.)
- True external sources (HTTP, Web Services, MQ, PowerExchange, etc.)

Public functions
----------------
1) extract_target_bound_fields(mapping_json, ...)
    - scopes to configured final target definition(s)
    - follows CONNECTOR links backward from each target field
    - walks pass-through ports and expression dependencies recursively
    - records lookup details and filter/router conditions on the traversed path
    - returns one lineage record per target field

2) humanise_rule(expression)
   - converts Informatica expressions into plain English

3) build_briefing_docx(mapping_json, ...)
   - builds a Word `.docx` briefing from extracted lineage

Compatibility alias
-------------------
`build_breifing_docx(...)` is provided as an alias for typo compatibility.
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
import tempfile
import textwrap
import time
from collections import defaultdict, deque
from datetime import datetime
from pathlib import Path
from typing import Any

# ============================================================
# LOGGING UTILITIES
# ============================================================

def _log(message: str, level: str = "INFO") -> None:
    """Print timestamped log message to terminal."""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    print(f"[{timestamp}] [{level:7s}] {message}")

def _log_info(message: str) -> None:
    _log(message, "INFO")

def _log_success(message: str) -> None:
    _log(message, "✓ OK")

def _log_warning(message: str) -> None:
    _log(message, "⚠ WARN")

def _log_error(message: str) -> None:
    _log(message, "✗ ERROR")

def _log_progress(current: int, total: int, message: str) -> None:
    _log(f"[{current}/{total}] {message}", "PROGRESS")

def _log_diagnostic(category: str, message: str) -> None:
    """Emit structured diagnostic warning (Strategy §8)."""
    _log(f"[DIAG:{category}] {message}", "⚠ DIAG")


# ============================================================
# DIAGNOSTICS COLLECTOR  (Strategy §8)
# ============================================================

class _Diagnostics:
    """
    Accumulates diagnostic warnings across the full extraction run.
    Categories per strategy doc: unresolved_expression, unresolved_port,
    missing_connector, recursive_structure, unsupported_construct.
    """
    def __init__(self) -> None:
        self._entries: list[dict[str, str]] = []

    def warn(self, category: str, mapping: str, instance: str, field: str, detail: str) -> None:
        entry = {
            "category": category,
            "mapping": mapping,
            "instance": instance,
            "field": field,
            "detail": detail,
        }
        self._entries.append(entry)
        _log_diagnostic(category, f"{mapping}/{instance}.{field} — {detail}")

    def all_entries(self) -> list[dict[str, str]]:
        return list(self._entries)

    def summary(self) -> dict[str, int]:
        counts: dict[str, int] = defaultdict(int)
        for e in self._entries:
            counts[e["category"]] += 1
        return dict(counts)


class IncompleteLineageError(RuntimeError):
    """
    Raised when strict completeness mode is enabled and one or more target
    fields could not be fully resolved (unresolved expressions, missing
    connectors, unvisited reachable nodes, or unsupported constructs).

    Carries a structured diagnostic report (see `_build_completeness_report`)
    so the caller can inspect, persist, or act on the failure without having
    to re-run extraction in best-effort mode.
    """

    def __init__(self, message: str, report: dict[str, Any]) -> None:
        super().__init__(message)
        self.report = report


# ============================================================
# TRANSFORMATION STRATEGY REGISTRY
# ============================================================

TRUE_EXTERNAL_SOURCES = {
    "SOURCE",
    "SOURCE DEFINITION",
    "SOURCE QUALIFIER",
    "ORACLE SOURCE QUALIFIER",
    "DB2 SOURCE QUALIFIER",
    "SQL SERVER SOURCE QUALIFIER",
    "TERADATA SOURCE QUALIFIER",
    "SAP SOURCE QUALIFIER",
    "SIEBEL SOURCE QUALIFIER",
    "PEOPLESOFT SOURCE QUALIFIER",
    "XML SOURCE QUALIFIER",
    "HTTP",
    "WEB SERVICES CONSUMER",
    "MQ SOURCE QUALIFIER",
    "JMS",
    "POWEREXCHANGE",
    "CDC SOURCE",
    "ODBC SOURCE QUALIFIER",
    "SEQUENCE",
}

PURE_PASSTHROUGH_TYPES = {
    "FILTER",
    "SORTER",
    "TRANSACTION CONTROL",
}

SPECIAL_PORT_TYPES = {
    "RANK": {"RANKINDEX"},
    "UPDATE STRATEGY": {"DD_UPDATE", "DD_INSERT", "DD_DELETE", "DD_REJECT"},
    "NORMALIZER": {"GK_", "GCID_"},
    "ROUTER": set(),
    "JOINER": set(),
    "UNION": set(),
}

OPAQUE_TYPES = {
    "JAVA TRANSFORMATION",
    "JAVA_TRANSFORMATION",
    "SQL TRANSFORMATION",
    "SQL_TRANSFORMATION",
    "CUSTOM TRANSFORMATION",
    "EXTERNAL PROCEDURE",
    "EXPRESSION",
}

MULTI_BRANCH_TYPES = {
    "UNION",
    "JOINER",
    "ROUTER",
    "DECODE",
}


# ============================================================
# GENERIC HELPERS
# ============================================================

def _as_list(value: Any) -> list[Any]:
    if value is None:
        return []
    if isinstance(value, list):
        return value
    return [value]


def _root_node(doc: dict[str, Any]) -> dict[str, Any]:
    """Return canonical root node from document or pass-through node."""
    if isinstance(doc, dict) and "tag" in doc and "attributes" in doc:
        return doc
    root = doc.get("root")
    if isinstance(root, dict) and "tag" in root:
        return root
    raise ValueError("Input JSON is not canonical Informatica node JSON.")


def _build_parent_map(root: dict[str, Any]) -> dict[int, dict[str, Any]]:
    """Build id(child) -> parent lookup for upward navigation."""
    parent_map: dict[int, dict[str, Any]] = {}
    stack = [root]
    while stack:
        node = stack.pop()
        for child in _as_list(node.get("children")):
            if isinstance(child, dict) and "tag" in child:
                parent_map[id(child)] = node
                stack.append(child)
    return parent_map


def _iter_nodes(root: dict[str, Any]):
    """Depth-first iteration over every canonical node."""
    stack = [root]
    while stack:
        node = stack.pop()
        yield node
        children = _as_list(node.get("children"))
        for child in reversed(children):
            if isinstance(child, dict) and "tag" in child:
                stack.append(child)


def _children_by_tag(node: dict[str, Any], tag: str) -> list[dict[str, Any]]:
    return [
        c for c in _as_list(node.get("children"))
        if isinstance(c, dict) and c.get("tag") == tag
    ]


def _ancestor(node: dict[str, Any], parent_map: dict[int, dict[str, Any]], tag: str) -> dict[str, Any] | None:
    cur = node
    while True:
        par = parent_map.get(id(cur))
        if par is None:
            return None
        if par.get("tag") == tag:
            return par
        cur = par


def _attr(node: dict[str, Any], key: str, default: str = "") -> str:
    attrs = node.get("attributes") or {}
    if not isinstance(attrs, dict):
        return default
    val = attrs.get(key, default)
    return "" if val is None else str(val)


# ============================================================
# WORKFLOW HELPERS
# ============================================================

def _workflow_terminal_session_mappings(root: dict[str, Any]) -> set[str]:
    """Return mapping names referenced by terminal session tasks in workflows."""
    session_name_to_mapping: dict[str, str] = {}
    for node in _iter_nodes(root):
        if node.get("tag") != "SESSION":
            continue
        s_name = _attr(node, "NAME")
        m_name = _attr(node, "MAPPINGNAME")
        if s_name and m_name:
            session_name_to_mapping[s_name] = m_name

    terminal_mappings: set[str] = set()
    for wf in (n for n in _iter_nodes(root) if n.get("tag") == "WORKFLOW"):
        session_tasks: set[str] = set()
        session_instance_names: set[str] = set()
        from_tasks: set[str] = set()

        for ti in _children_by_tag(wf, "TASKINSTANCE"):
            tasktype = _attr(ti, "TASKTYPE").upper()
            is_enabled = _attr(ti, "ISENABLED", "YES").upper()
            if tasktype != "SESSION" or is_enabled != "YES":
                continue
            tname = _attr(ti, "TASKNAME")
            iname = _attr(ti, "NAME")
            if tname:
                session_tasks.add(tname)
            if iname:
                session_instance_names.add(iname)

        for wl in _children_by_tag(wf, "WORKFLOWLINK"):
            from_task = _attr(wl, "FROMTASK")
            if from_task:
                from_tasks.add(from_task)

        terminal_tasks = session_tasks - from_tasks
        terminal_instances = session_instance_names - from_tasks
        for task_name in terminal_tasks | terminal_instances:
            mapping = session_name_to_mapping.get(task_name)
            if mapping:
                terminal_mappings.add(mapping)

    return terminal_mappings


def _mapping_names_in_order(root: dict[str, Any]) -> list[str]:
    """Return mapping names in XML order."""
    names: list[str] = []
    seen: set[str] = set()
    for node in _iter_nodes(root):
        if node.get("tag") != "MAPPING":
            continue
        name = _attr(node, "NAME")
        if name and name not in seen:
            seen.add(name)
            names.append(name)
    return names


# ============================================================
# INSTANCE TYPE HELPERS
# ============================================================

def _is_source_instance(instance_attrs: dict[str, str]) -> bool:
    t = str(instance_attrs.get("TYPE", "")).upper()
    tt = str(instance_attrs.get("TRANSFORMATION_TYPE", "")).upper()
    return (
        t == "SOURCE"
        or "SOURCE DEFINITION" in tt
        or "SOURCE QUALIFIER" in tt
    )


def _is_target_instance(instance_attrs: dict[str, str]) -> bool:
    t = str(instance_attrs.get("TYPE", "")).upper()
    tt = str(instance_attrs.get("TRANSFORMATION_TYPE", "")).upper()
    return t == "TARGET" or "TARGET DEFINITION" in tt


def _is_transformation_instance(instance_attrs: dict[str, str]) -> bool:
    if _is_source_instance(instance_attrs) or _is_target_instance(instance_attrs):
        return False
    t = str(instance_attrs.get("TYPE", "")).upper()
    tt = str(instance_attrs.get("TRANSFORMATION_TYPE", "")).upper()
    return t == "TRANSFORMATION" or bool(tt)


def _tx_type(instance_attrs: dict[str, str]) -> str:
    return str(instance_attrs.get("TRANSFORMATION_TYPE") or instance_attrs.get("TYPE") or "")


def _is_lookup_type(tx_type: str) -> bool:
    return "LOOKUP" in tx_type.upper()


def _is_filter_or_router_type(tx_type: str) -> bool:
    up = tx_type.upper()
    return "FILTER" in up or "ROUTER" in up


# ============================================================
# EXPRESSION UTILITIES
# ============================================================

def _is_pass_through_expression(expression: str, field_name: str) -> bool:
    expr = (expression or "").strip()
    if not expr:
        return True
    if re.fullmatch(r"[A-Za-z_][A-Za-z0-9_$#]*", expr):
        return True
    if field_name and expr.upper() == field_name.upper():
        return True
    return False


def _expression_port_dependencies(expression: str, valid_port_names: list[str]) -> list[str]:
    """Return referenced local port names from an expression, preserving order."""
    expr = (expression or "").strip()
    if not expr:
        return []
    by_upper = {name.upper(): name for name in valid_port_names}
    deps: list[str] = []
    seen: set[str] = set()
    for token in re.findall(r"[A-Za-z_][A-Za-z0-9_$#]*", expr):
        key = token.upper()
        if key in by_upper and key not in seen:
            seen.add(key)
            deps.append(by_upper[key])
    return deps


def _lookup_condition_pairs(lookup_condition: str) -> list[tuple[str, str]]:
    """Extract identifier pairs from lookup conditions like A = B AND C = D."""
    condition = (lookup_condition or "").strip()
    if not condition:
        return []
    pairs: list[tuple[str, str]] = []
    seen: set[tuple[str, str]] = set()
    clauses = re.split(r"\bAND\b|\bOR\b", condition, flags=re.IGNORECASE)
    for clause in clauses:
        part = clause.strip().strip("()")
        if not part or "=" not in part:
            continue
        m = re.search(r"(.+?)\s*=\s*(.+)", part)
        if not m:
            continue
        left_tokens = re.findall(r"[A-Za-z_][A-Za-z0-9_$#.]*", m.group(1))
        right_tokens = re.findall(r"[A-Za-z_][A-Za-z0-9_$#.]*", m.group(2))
        if not left_tokens or not right_tokens:
            continue
        left = left_tokens[-1].split(".")[-1]
        right = right_tokens[-1].split(".")[-1]
        key = (left.upper(), right.upper())
        if key not in seen:
            seen.add(key)
            pairs.append((left, right))
    return pairs


def _extract_opaque_field_mapping(expression: str, output_field: str) -> str | None:
    """
    Extract input field reference from opaque transformation expressions.
    E.g., 'REFITKEYID = I_REFITKEYID' -> 'I_REFITKEYID'
    """
    expr = (expression or "").strip()
    if not expr or "=" not in expr:
        return None
    m = re.match(r"^\s*([A-Za-z_][A-Za-z0-9_$#]*)\s*=\s*([A-Za-z_][A-Za-z0-9_$#]*)\s*$", expr)
    if not m:
        return None
    left = m.group(1)
    right = m.group(2)
    if left.upper() != output_field.upper():
        return None
    return right


def _split_sql_select_items(select_sql: str) -> list[str]:
    """Split SELECT-list by commas while respecting parentheses and quotes."""
    items: list[str] = []
    cur: list[str] = []
    depth = 0
    in_single = False
    in_double = False
    i = 0
    while i < len(select_sql):
        ch = select_sql[i]
        if ch == "'" and not in_double:
            in_single = not in_single
            cur.append(ch)
            i += 1
            continue
        if ch == '"' and not in_single:
            in_double = not in_double
            cur.append(ch)
            i += 1
            continue
        if not in_single and not in_double:
            if ch == "(":
                depth += 1
            elif ch == ")" and depth > 0:
                depth -= 1
            elif ch == "," and depth == 0:
                item = "".join(cur).strip()
                if item:
                    items.append(item)
                cur = []
                i += 1
                continue
        cur.append(ch)
        i += 1
    tail = "".join(cur).strip()
    if tail:
        items.append(tail)
    return items


def _sql_alias_derivation(sql_override: str, alias_name: str) -> dict[str, Any]:
    sql = (sql_override or "").strip()
    alias = (alias_name or "").strip()
    if not sql or not alias:
        return {}
    expr = ""
    m = re.search(r"(?is)\bselect\b\s+(.*?)\bfrom\b", sql)
    if m:
        select_list = m.group(1)
        for item in _split_sql_select_items(select_list):
            item_norm = re.sub(r"\s+", " ", item).strip()
            m_as = re.search(r"(?is)^(.+?)\s+as\s+([A-Za-z_][A-Za-z0-9_$#]*)\s*$", item_norm)
            if m_as and m_as.group(2).upper() == alias.upper():
                expr = m_as.group(1).strip()
                break
            m_alias = re.search(r"(?is)^(.+?)\s+([A-Za-z_][A-Za-z0-9_$#]*)\s*$", item_norm)
            if m_alias and m_alias.group(2).upper() == alias.upper():
                expr = m_alias.group(1).strip()
                break
            simple_name = item_norm.split(".")[-1]
            if re.fullmatch(r"[A-Za-z_][A-Za-z0-9_$#]*", simple_name) and simple_name.upper() == alias.upper():
                expr = item_norm
                break
    if not expr:
        m_fallback = re.search(
            rf"(?is)([^,\n\r]+?)\s+as\s+{re.escape(alias)}\b",
            sql,
        )
        if m_fallback:
            expr = m_fallback.group(1).strip()
    if not expr:
        return {}
    expr = re.sub(r"\s+", " ", expr)
    cols = []
    seen = set()
    for c in re.findall(r"[A-Za-z_][A-Za-z0-9_$#]*\.[A-Za-z_][A-Za-z0-9_$#]*", expr):
        key = c.upper()
        if key not in seen:
            seen.add(key)
            cols.append(c)
    return {
        "lookup_sql_alias": alias,
        "lookup_sql_expression": expr,
        "lookup_sql_source_columns": cols,
    }


def _is_condition_or_query_attribute(attr_name: str) -> bool:
    n = (attr_name or "").strip().lower()
    if not n:
        return False
    keys = ("condition", "filter", "sql", "query", "override", "join", "where", "having")
    return any(k in n for k in keys)


def _is_generated_port(tx_type: str, port_name: str) -> bool:
    tx_type_upper = tx_type.upper()
    if port_name == "RANKINDEX" and "RANK" in tx_type_upper:
        return True
    if port_name == "TCONTROL" and "TRANSACTION CONTROL" in tx_type_upper:
        return True
    if port_name.startswith("DD_") and "UPDATE STRATEGY" in tx_type_upper:
        return True
    if (port_name.startswith("GK_") or port_name.startswith("GCID_")) and "NORMALIZER" in tx_type_upper:
        return True
    return False


def _get_base_instance_name(instance_name: str, tx_type: str) -> str:
    """Strip dot-suffix from Router/Joiner instance names to get base name."""
    tx_upper = tx_type.upper()
    if ("ROUTER" in tx_upper or "JOINER" in tx_upper) and "." in instance_name:
        return instance_name.split(".")[0]
    return instance_name


def _extract_unconnected_lookup_refs(expression: str) -> list[tuple[str, str]]:
    refs = []
    pattern = r'LKP\s*:\s*([A-Za-z_][A-Za-z0-9_]*)\s*\(\s*([A-Za-z_][A-Za-z0-9_]*)\s*\)'
    for m in re.finditer(pattern, expression, re.IGNORECASE):
        refs.append((m.group(1), m.group(2)))
    return refs


def _extract_sql_bind_variables(sql_query: str) -> list[str]:
    vars_: list[str] = []
    seen: set[str] = set()
    for m in re.finditer(r':([A-Za-z_][A-Za-z0-9_]*)', sql_query):
        var = m.group(1)
        if var.upper() not in seen:
            seen.add(var.upper())
            vars_.append(var)
    return vars_


# ============================================================
# MAPPING INDEX
# ============================================================

class _MappingIndex:
    def __init__(self, mapping_node: dict[str, Any], folder_node: dict[str, Any]):
        self.mapping_node = mapping_node
        self.folder_node = folder_node
        self.mapping_name = _attr(mapping_node, "NAME", "<unknown_mapping>")
        self.folder_name = _attr(folder_node, "NAME", "<unknown_folder>")

        self.instances: dict[str, dict[str, str]] = {}
        for inst in _children_by_tag(mapping_node, "INSTANCE"):
            attrs = inst.get("attributes") or {}
            name = str(attrs.get("NAME", ""))
            if name:
                self.instances[name] = {k: "" if v is None else str(v) for k, v in attrs.items()}

        self.transformations: dict[str, dict[str, Any]] = {}
        self.transform_field_by_name: dict[str, dict[str, dict[str, str]]] = defaultdict(dict)
        self.transform_table_attributes: dict[str, dict[str, str]] = defaultdict(dict)

        def _index_transformation(tx: dict[str, Any]) -> None:
            tx_name = _attr(tx, "NAME")
            if not tx_name:
                return
            # Only store mapping-local tx in self.transformations (used for type lookups
            # against mapping instances); folder-level reusable transformations are only
            # indexed into the field/attribute tables below so the metadata is available.
            if tx_name not in self.transformations:
                self.transformations[tx_name] = tx
            for tf in _children_by_tag(tx, "TRANSFORMFIELD"):
                tf_name = _attr(tf, "NAME")
                if tf_name:
                    tf_attrs = tf.get("attributes") or {}
                    self.transform_field_by_name[tx_name][tf_name] = {
                        k: "" if v is None else str(v)
                        for k, v in tf_attrs.items()
                    }
            for ta in _children_by_tag(tx, "TABLEATTRIBUTE"):
                ta_name = _attr(ta, "NAME")
                if not ta_name:
                    continue
                self.transform_table_attributes[tx_name][ta_name] = _attr(ta, "VALUE")

        # Index folder-level reusable transformations FIRST so that mapping-local
        # definitions (if any) can override them in the second pass.
        for tx in _children_by_tag(folder_node, "TRANSFORMATION"):
            _index_transformation(tx)

        # Index mapping-local (embedded) transformations — these take precedence.
        for tx in _children_by_tag(mapping_node, "TRANSFORMATION"):
            tx_name = _attr(tx, "NAME")
            if tx_name:
                # Overwrite any folder-level entry so local wins.
                self.transformations[tx_name] = tx
            _index_transformation(tx)

        # Fix 2: Merge instance-level TABLEATTRIBUTE overrides.
        # In Informatica XML an INSTANCE element can carry its own TABLEATTRIBUTE
        # children (e.g. "Lookup Sql Override") that shadow or supplement the
        # reusable transformation's attributes.  Index them here so that
        # _trace_lineage_from_node picks them up via transform_table_attributes.
        for inst in _children_by_tag(mapping_node, "INSTANCE"):
            inst_attrs = inst.get("attributes") or {}
            # The logical name of this instance is NAME; the underlying transformation
            # definition name is TRANSFORMATION_NAME (may differ for reusable objects).
            inst_name = str(inst_attrs.get("NAME", ""))
            tx_def_name = str(inst_attrs.get("TRANSFORMATION_NAME", "")) or inst_name
            for ta in _children_by_tag(inst, "TABLEATTRIBUTE"):
                ta_name = _attr(ta, "NAME")
                if not ta_name:
                    continue
                # Write under BOTH the instance name and the definition name so that
                # lookups via either key always succeed.
                val = _attr(ta, "VALUE")
                if inst_name:
                    self.transform_table_attributes[inst_name][ta_name] = val
                if tx_def_name and tx_def_name != inst_name:
                    self.transform_table_attributes[tx_def_name][ta_name] = val
        self.backward_edges: dict[tuple[str, str], list[tuple[str, str]]] = defaultdict(list)
        self.forward_edges: dict[tuple[str, str], list[tuple[str, str]]] = defaultdict(list)
        self.incoming_target_fields: dict[str, set[str]] = defaultdict(set)
        self.all_connectors: list[dict[str, str]] = []

        for conn in _children_by_tag(mapping_node, "CONNECTOR"):
            attrs = conn.get("attributes") or {}
            from_inst = str(attrs.get("FROMINSTANCE", ""))
            from_field = str(attrs.get("FROMFIELD", ""))
            to_inst = str(attrs.get("TOINSTANCE", ""))
            to_field = str(attrs.get("TOFIELD", ""))
            if not (from_inst and from_field and to_inst and to_field):
                continue
            self.backward_edges[(to_inst, to_field)].append((from_inst, from_field))
            self.forward_edges[(from_inst, from_field)].append((to_inst, to_field))
            if to_inst:
                self.incoming_target_fields[to_inst].add(to_field)
            self.all_connectors.append({
                "from_instance": from_inst,
                "from_field": from_field,
                "to_instance": to_inst,
                "to_field": to_field,
            })

        # Secondary index for Union-type traversal: to_field_upper -> list of (to_inst, to_field) keys.
        # Avoids a full O(N) backward_edges scan when collecting all upstream Union group connectors.
        # Built once here; the Union branch in _trace_lineage_from_node uses it for O(k) lookup
        # where k = number of connectors sharing that field name (typically very small).
        self.backward_keys_by_to_field: dict[str, list[tuple[str, str]]] = defaultdict(list)
        for (to_inst, to_field) in self.backward_edges:
            self.backward_keys_by_to_field[to_field.upper()].append((to_inst, to_field))

        self.target_fields_by_definition: dict[str, list[str]] = {}
        target_defs_upper: set[str] = set()
        for tgt in _children_by_tag(folder_node, "TARGET"):
            tgt_name = _attr(tgt, "NAME")
            if not tgt_name:
                continue
            target_defs_upper.add(tgt_name.upper())
            fields = []
            for tf in _children_by_tag(tgt, "TARGETFIELD"):
                name = _attr(tf, "NAME")
                if name:
                    fields.append(name)
            self.target_fields_by_definition[tgt_name] = fields

        self.source_definition_names_upper: set[str] = set()
        for src in _children_by_tag(folder_node, "SOURCE"):
            src_name = _attr(src, "NAME")
            if src_name:
                self.source_definition_names_upper.add(src_name.upper())

        self.final_target_definition_names_upper: set[str] = (
            target_defs_upper - self.source_definition_names_upper
        )

        self.target_load_order: list[tuple[int, str]] = []
        for tlo in _children_by_tag(mapping_node, "TARGETLOADORDER"):
            attrs = tlo.get("attributes") or {}
            target_instance = str(attrs.get("TARGETINSTANCE", ""))
            try:
                order = int(str(attrs.get("ORDER", "")).strip())
            except ValueError:
                order = -1
            if target_instance:
                self.target_load_order.append((order, target_instance))

    def transformation_name_for_instance(self, instance_name: str) -> str:
        attrs = self.instances.get(instance_name, {})
        tx_name = attrs.get("TRANSFORMATION_NAME", "")
        return tx_name or instance_name

    def all_lookup_instances(self) -> list[dict[str, str]]:
        """
        Fix 4: Return metadata for every lookup instance in the mapping,
        regardless of whether backward traversal reached it.

        This ensures that lookups behind custom Union Transformation nodes (or
        any other opaque/unrecognised branching type) are never silently dropped
        from the output.  The caller merges these into lookup_used after the
        main traversal, skipping duplicates by instance name.
        """
        results: list[dict[str, str]] = []
        for inst_name, inst_attrs in self.instances.items():
            if not _is_transformation_instance(inst_attrs):
                continue
            tx_type = _tx_type(inst_attrs)
            if not _is_lookup_type(tx_type):
                continue
            tx_name = self.transformation_name_for_instance(inst_name)
            # Merge table attributes: definition-level first, then instance-level
            # overrides (instance-level was indexed under inst_name in Fix 2).
            ta: dict[str, str] = {}
            ta.update(self.transform_table_attributes.get(tx_name, {}))
            if inst_name != tx_name:
                ta.update(self.transform_table_attributes.get(inst_name, {}))
            results.append({
                "instance_name": inst_name,
                "transform_name": tx_name,
                "lookup_table": ta.get("Lookup table name", ""),
                "lookup_condition": ta.get("Lookup condition", ""),
                "lookup_sql_override": ta.get("Lookup Sql Override", ""),
                "lookup_source_filter": ta.get("Lookup Source Filter", ""),
            })
        return results

    def all_reachable_connectors_from_targets(
        self, target_instance_names: list[str]
    ) -> set[tuple[str, str]]:
        """
        Forward pass to pre-compute all connectors reachable from sources
        that eventually reach any of the given target instances.
        Used by completeness validation. (Strategy §7)
        """
        # BFS backward from all target ports
        reachable: set[tuple[str, str]] = set()
        queue: deque[tuple[str, str]] = deque()
        for ti in target_instance_names:
            for field in self.incoming_target_fields.get(ti, set()):
                queue.append((ti, field))
        while queue:
            node = queue.popleft()
            if node in reachable:
                continue
            reachable.add(node)
            for parent in self.backward_edges.get(node, []):
                if parent not in reachable:
                    queue.append(parent)
        return reachable


# ============================================================
# CHECKPOINT / PERSISTENCE HELPERS  (Strategy §2, §3)
# ============================================================

class _CheckpointManager:
    """
    Persists intermediate processing artifacts to a temp directory.
    Each stage saves its output; if processing is interrupted the last
    checkpoint can be reloaded without recomputation. (Strategy §2, §3)
    """

    STAGES = [
        "parsed_metadata",      # transformation metadata per mapping
        "connector_graph",      # backward/forward edge tables
        "port_dependency_graph",# expression-level port deps per tx
        "expression_dep_map",   # full expression dependency map
        "lookup_metadata",      # lookup table/condition data
        "router_group_metadata",# router group assignments
        "joiner_metadata",      # join metadata
        "aggregator_metadata",  # aggregator metadata
        "partial_lineage_cache",# memoised lineage results
        "traversal_checkpoints",# per-mapping traversal state
    ]

    def __init__(self, base_dir: str | Path | None = None) -> None:
        if base_dir:
            self._dir = Path(base_dir)
            self._dir.mkdir(parents=True, exist_ok=True)
            self._temp_dir_obj = None
        else:
            self._temp_dir_obj = tempfile.TemporaryDirectory(prefix="infa_lineage_")
            self._dir = Path(self._temp_dir_obj.name)
        _log_info(f"Checkpoint directory: {self._dir}")

    def save(self, stage: str, mapping_name: str, data: Any) -> None:
        safe_name = re.sub(r"[^\w\-]", "_", mapping_name)
        path = self._dir / f"{stage}__{safe_name}.json"
        try:
            path.write_text(
                json.dumps(data, indent=2, ensure_ascii=False, default=str),
                encoding="utf-8",
            )
        except Exception as e:
            _log_warning(f"Checkpoint save failed [{stage}][{mapping_name}]: {e}")

    def load(self, stage: str, mapping_name: str) -> Any | None:
        safe_name = re.sub(r"[^\w\-]", "_", mapping_name)
        path = self._dir / f"{stage}__{safe_name}.json"
        if not path.exists():
            return None
        try:
            return json.loads(path.read_text(encoding="utf-8"))
        except Exception as e:
            _log_warning(f"Checkpoint load failed [{stage}][{mapping_name}]: {e}")
            return None

    def validate_counts(
        self, stage: str, mapping_name: str, expected: int, actual: int
    ) -> bool:
        """Validate count consistency after a stage. (Strategy §3)"""
        if expected != actual:
            _log_warning(
                f"Count mismatch at [{stage}] for {mapping_name}: "
                f"expected {expected}, got {actual}"
            )
            return False
        return True

    def cleanup(self) -> None:
        if self._temp_dir_obj:
            try:
                self._temp_dir_obj.cleanup()
            except Exception:
                pass


# ============================================================
# GRAPH MEMOIZATION CACHE  (Strategy §4)
# ============================================================

class _LineageCache:
    """
    Memoises lineage traversal results keyed by (instance, port).
    Prevents redundant re-traversals across multiple target fields
    that share upstream paths. (Strategy §4)
    """

    def __init__(self) -> None:
        self._cache: dict[tuple[str, str], dict[str, Any]] = {}
        self._hits = 0
        self._misses = 0

    def get(self, node: tuple[str, str]) -> dict[str, Any] | None:
        result = self._cache.get(node)
        if result is not None:
            self._hits += 1
        else:
            self._misses += 1
        return result

    def put(self, node: tuple[str, str], result: dict[str, Any]) -> None:
        self._cache[node] = result

    def stats(self) -> dict[str, int]:
        return {"hits": self._hits, "misses": self._misses, "size": len(self._cache)}

    def clear(self) -> None:
        self._cache.clear()
        self._hits = 0
        self._misses = 0


# ============================================================
# CORE LINEAGE TRAVERSAL  (Strategy §5, §6)
# ============================================================

def _trace_lineage_from_node(
    idx: _MappingIndex,
    start_node: tuple[str, str],
    lineage_cache: _LineageCache | None = None,
    diagnostics: _Diagnostics | None = None,
    _cycle_stack: frozenset[tuple[str, str]] | None = None,
) -> dict[str, Any]:
    """
    Trace lineage backwards from any node (instance, field).

    Key improvements over v1 (Strategy §5, §6):
    - Path-aware visited tracking: separates cycle_stack (active DFS path)
      from visited (completed nodes). A node can be revisited on a different
      path — only active-path cycles are suppressed, not valid branches.
    - Graph memoization: results keyed by (instance, port) are reused.
    - Robust traversal: opaque transformations get name-match fallback
      instead of terminating the traversal.
    - Diagnostics: unresolved ports, missing connectors, recursive
      structures are flagged rather than silently skipped.
    """
    if _cycle_stack is None:
        _cycle_stack = frozenset()

    # --- Memoization check (Strategy §4) ---
    if lineage_cache is not None:
        cached = lineage_cache.get(start_node)
        if cached is not None:
            return cached

    queue: deque[tuple[tuple[str, str], frozenset[tuple[str, str]]]] = deque(
        [(start_node, _cycle_stack)]
    )

    # Separate completed-node set from active cycle detection (Strategy §6)
    completed: set[tuple[str, str]] = set()
    visit_order: list[tuple[str, str]] = []
    source_nodes: list[tuple[str, str]] = []
    source_seen: set[tuple[str, str]] = set()
    touched: list[dict[str, Any]] = []
    touched_seen: set[tuple[str, str]] = set()
    lookup_used: list[dict[str, Any]] = []
    lookup_seen: set[str] = set()
    conditions_applied: list[dict[str, str]] = []
    condition_seen: set[tuple[str, str]] = set()
    stop_expressions: list[dict[str, str]] = []
    unresolved_nodes: list[dict[str, str]] = []  # diagnostics

    while queue:
        node, active_path = queue.popleft()

        # Completed-node cache — don't re-process (Strategy §6)
        if node in completed:
            continue

        # Cycle detection on active path only (Strategy §6)
        if node in active_path:
            if diagnostics:
                inst_n, field_n = node
                diagnostics.warn(
                    "recursive_structure",
                    idx.mapping_name,
                    inst_n,
                    field_n,
                    f"Cycle detected on active traversal path — skipping branch",
                )
            continue

        completed.add(node)
        visit_order.append(node)
        new_active_path = active_path | {node}

        inst_name, field_name = node
        inst_attrs = idx.instances.get(inst_name, {})

        # Validate the instance exists in this mapping
        if not inst_attrs and inst_name:
            if diagnostics:
                diagnostics.warn(
                    "unresolved_port",
                    idx.mapping_name,
                    inst_name,
                    field_name,
                    "Instance not found in mapping index",
                )
            unresolved_nodes.append({"instance": inst_name, "field": field_name, "reason": "unknown_instance"})
            continue

        should_walk_parents = True
        internal_dependency_nodes: list[tuple[str, str]] = []

        if _is_transformation_instance(inst_attrs):
            tx_name = idx.transformation_name_for_instance(inst_name)
            tx_fields = idx.transform_field_by_name.get(tx_name, {})
            tf_attrs = tx_fields.get(field_name, {})
            tx_field_names = list(tx_fields.keys())
            tx_field_name_by_upper = {n.upper(): n for n in tx_field_names}
            tx_type = _tx_type(inst_attrs)
            expr = tf_attrs.get("EXPRESSION", "")
            porttype = tf_attrs.get("PORTTYPE", "").upper()

            # Validate field exists in transformation (Strategy §8)
            if not tf_attrs and field_name:
                # Try case-insensitive match before giving up
                canonical = tx_field_name_by_upper.get(field_name.upper())
                if canonical:
                    field_name = canonical
                    tf_attrs = tx_fields.get(field_name, {})
                    expr = tf_attrs.get("EXPRESSION", "")
                    porttype = tf_attrs.get("PORTTYPE", "").upper()
                    node = (inst_name, field_name)
                else:
                    if diagnostics:
                        diagnostics.warn(
                            "unresolved_port",
                            idx.mapping_name,
                            inst_name,
                            field_name,
                            f"Port '{field_name}' not found in transformation '{tx_name}'",
                        )
                    unresolved_nodes.append({
                        "instance": inst_name,
                        "field": field_name,
                        "reason": "port_not_in_transformation",
                        "transformation": tx_name,
                    })

            # --- Generated ports (RANKINDEX, DD_*, TCONTROL, GK_, GCID_) ---
            if _is_generated_port(tx_type, field_name):
                key = (inst_name, field_name)
                if key not in touched_seen:
                    touched_seen.add(key)
                    touched.append({
                        "instance_name": inst_name,
                        "transform_name": tx_name,
                        "transform_type": tx_type,
                        "field_name": field_name,
                        "expression": f"[GENERATED: {tx_type}]",
                        "porttype": porttype,
                        "datatype": tf_attrs.get("DATATYPE", ""),
                        "precision": tf_attrs.get("PRECISION", ""),
                        "scale": tf_attrs.get("SCALE", ""),
                    })
                continue

            # --- Return ports (lookup condition-driven) ---
            if porttype == "RETURN" and _is_lookup_type(tx_type):
                key = (inst_name, field_name)
                if key not in touched_seen:
                    touched_seen.add(key)
                    touched.append({
                        "instance_name": inst_name,
                        "transform_name": tx_name,
                        "transform_type": tx_type,
                        "field_name": field_name,
                        "expression": "[RETURN: lookup result]",
                        "porttype": porttype,
                        "datatype": tf_attrs.get("DATATYPE", ""),
                        "precision": tf_attrs.get("PRECISION", ""),
                        "scale": tf_attrs.get("SCALE", ""),
                    })
                if _is_source_instance(inst_attrs):
                    if node not in source_seen:
                        source_seen.add(node)
                        source_nodes.append(node)
                continue

            # --- Internal expression dependencies ---
            for dep_field in _expression_port_dependencies(expr, tx_field_names):
                if dep_field != field_name:
                    internal_dependency_nodes.append((inst_name, dep_field))

            key = (inst_name, field_name)
            if key not in touched_seen:
                touched_seen.add(key)
                touched.append({
                    "instance_name": inst_name,
                    "transform_name": tx_name,
                    "transform_type": tx_type,
                    "field_name": field_name,
                    "expression": expr,
                    "porttype": porttype,
                    "datatype": tf_attrs.get("DATATYPE", ""),
                    "precision": tf_attrs.get("PRECISION", ""),
                    "scale": tf_attrs.get("SCALE", ""),
                })

            # --- Lookup transformations ---
            if _is_lookup_type(tx_type):
                ta = idx.transform_table_attributes.get(tx_name, {})
                lk_cond = ta.get("Lookup condition", "")
                for left_name, right_name in _lookup_condition_pairs(lk_cond):
                    left_field = tx_field_name_by_upper.get(left_name.upper())
                    right_field = tx_field_name_by_upper.get(right_name.upper())
                    if not left_field or not right_field:
                        if diagnostics:
                            missing = left_name if not left_field else right_name
                            diagnostics.warn(
                                "unresolved_expression",
                                idx.mapping_name,
                                inst_name,
                                field_name,
                                f"Lookup condition references unknown port '{missing}'",
                            )
                        continue
                    if left_field.upper() == field_name.upper() and right_field.upper() != left_field.upper():
                        internal_dependency_nodes.append((inst_name, right_field))
                    elif right_field.upper() == field_name.upper() and left_field.upper() != right_field.upper():
                        internal_dependency_nodes.append((inst_name, left_field))

                if tx_name not in lookup_seen:
                    lookup_seen.add(tx_name)
                    lookup_used.append({
                        "instance_name": inst_name,
                        "transform_name": tx_name,
                        "lookup_table": ta.get("Lookup table name", ""),
                        "lookup_condition": lk_cond,
                        "lookup_sql_override": ta.get("Lookup Sql Override", ""),
                        "lookup_source_filter": ta.get("Lookup Source Filter", ""),
                    })

                # Handle unconnected lookups referenced in expressions (Strategy §5)
                if expr:
                    for lkp_name, lkp_port in _extract_unconnected_lookup_refs(expr):
                        if lkp_name not in lookup_seen:
                            lookup_seen.add(lkp_name)
                            lkp_tx_name = idx.transformation_name_for_instance(lkp_name) if lkp_name in idx.instances else lkp_name
                            lkp_ta = idx.transform_table_attributes.get(lkp_tx_name, {})
                            lookup_used.append({
                                "instance_name": lkp_name,
                                "transform_name": lkp_tx_name,
                                "lookup_table": lkp_ta.get("Lookup table name", ""),
                                "lookup_condition": lkp_ta.get("Lookup condition", ""),
                                "lookup_sql_override": lkp_ta.get("Lookup Sql Override", ""),
                                "lookup_source_filter": lkp_ta.get("Lookup Source Filter", ""),
                                "unconnected": True,
                                "referenced_port": lkp_port,
                            })

            # --- Filter / Router conditions ---
            if _is_filter_or_router_type(tx_type):
                ta = idx.transform_table_attributes.get(tx_name, {})
                for k, v in ta.items():
                    if "condition" in k.lower():
                        ckey = (tx_name, k)
                        if ckey not in condition_seen:
                            condition_seen.add(ckey)
                            conditions_applied.append({
                                "instance_name": inst_name,
                                "transform_name": tx_name,
                                "condition_name": k,
                                "condition_value": v,
                            })

            # --- SQL bind variables (for SQL Transformation) ---
            if "SQL" in tx_type.upper():
                sql_query = idx.transform_table_attributes.get(tx_name, {}).get("Sql Query", "")
                if sql_query:
                    for bind_var in _extract_sql_bind_variables(sql_query):
                        canon = tx_field_name_by_upper.get(bind_var.upper())
                        if canon and canon != field_name:
                            internal_dependency_nodes.append((inst_name, canon))

            # --- Stop expression vs continue (Strategy §5: never terminate for opaque) ---
            opaque_input_field = None
            tx_type_up = tx_type.upper()
            if "JAVA" in tx_type_up or "SQL" in tx_type_up or "CUSTOM" in tx_type_up or "EXTERNAL" in tx_type_up:
                # Attempt to extract input field mapping from expression
                opaque_input_field = _extract_opaque_field_mapping(expr, field_name)

                if opaque_input_field:
                    canon = tx_field_name_by_upper.get(opaque_input_field.upper())
                    if canon:
                        internal_dependency_nodes.append((inst_name, canon))
                    else:
                        # Strategy §5: continue traversal via connectors even when opaque mapping fails
                        if diagnostics:
                            diagnostics.warn(
                                "unsupported_construct",
                                idx.mapping_name,
                                inst_name,
                                field_name,
                                f"Opaque type '{tx_type}' references unknown input port '{opaque_input_field}' — continuing via connectors",
                            )
                    if expr:
                        stop_expressions.append({
                            "instance_name": inst_name,
                            "transform_name": tx_name,
                            "field_name": field_name,
                            "expression": expr,
                        })
                else:
                    # No parseable mapping — still walk connectors (Strategy §5: robust traversal)
                    if expr and not _is_pass_through_expression(expr, field_name):
                        stop_expressions.append({
                            "instance_name": inst_name,
                            "transform_name": tx_name,
                            "field_name": field_name,
                            "expression": expr,
                        })
                        if diagnostics:
                            diagnostics.warn(
                                "unresolved_expression",
                                idx.mapping_name,
                                inst_name,
                                field_name,
                                f"Opaque '{tx_type}': expression not parseable as simple field mapping — traversal continues via connectors",
                            )
                    # Walk parents regardless for opaque types (Strategy §5)
                    should_walk_parents = True
            else:
                if expr and not _is_pass_through_expression(expr, field_name):
                    should_walk_parents = False
                    stop_expressions.append({
                        "instance_name": inst_name,
                        "transform_name": tx_name,
                        "field_name": field_name,
                        "expression": expr,
                    })

        connector_parents: list[tuple[str, str]] = []
        base_inst_name = _get_base_instance_name(inst_name, _tx_type(inst_attrs))
        inst_tx_type = _tx_type(inst_attrs).upper()

        # Fix 3: Detect "Custom Transformation" instances whose TEMPLATENAME
        # (stored on the underlying reusable TRANSFORMATION node) is
        # "Union Transformation" so they are handled identically to a native
        # UNION type -- i.e. ALL upstream group connectors are followed.
        def _is_union_type(i_name: str, i_tx_type: str) -> bool:
            if "UNION" in i_tx_type:
                return True
            if "CUSTOM" in i_tx_type:
                tx_def_name = idx.transformation_name_for_instance(i_name)
                tx_node = idx.transformations.get(tx_def_name)
                if tx_node is not None:
                    template_name = _attr(tx_node, "TEMPLATENAME", "")
                    if "UNION" in template_name.upper():
                        return True
            return False

        if base_inst_name != inst_name and ("ROUTER" in inst_tx_type or "JOINER" in inst_tx_type):
            connector_parents.extend(idx.backward_edges.get((base_inst_name, field_name), []))
        else:
            connector_parents.extend(idx.backward_edges.get((inst_name, field_name), []))

        # Union: collect ALL upstream group connectors for this field (Strategy 5).
        # Fix 3: also catches custom Union Transformation instances.
        # Optimised: use backward_keys_by_to_field secondary index (built in _MappingIndex)
        # so only connectors sharing the same to_field name are iterated — O(k) not O(N).
        if _is_union_type(inst_name, inst_tx_type):
            connector_parents = []
            _union_prefix = inst_name + "."
            for (to_inst, to_field) in idx.backward_keys_by_to_field.get(field_name.upper(), []):
                if to_inst == inst_name or to_inst.startswith(_union_prefix):
                    connector_parents.extend(idx.backward_edges[(to_inst, to_field)])

        # Diagnostic: flag if no connector found and not a source (Strategy §8)
        if (
            not connector_parents
            and not internal_dependency_nodes
            and not _is_source_instance(inst_attrs)
            and not _is_target_instance(inst_attrs)
            and _is_transformation_instance(inst_attrs)
            and should_walk_parents
        ):
            if diagnostics:
                diagnostics.warn(
                    "missing_connector",
                    idx.mapping_name,
                    inst_name,
                    field_name,
                    f"No incoming connectors found for transformation port — may indicate incomplete extraction",
                )
            unresolved_nodes.append({
                "instance": inst_name,
                "field": field_name,
                "reason": "no_incoming_connector",
            })

        # Assemble all parents to explore
        parents_to_walk: list[tuple[str, str]] = []
        if should_walk_parents:
            parents_to_walk.extend(connector_parents)
        parents_to_walk.extend(internal_dependency_nodes)

        if not parents_to_walk:
            if _is_source_instance(inst_attrs):
                if node not in source_seen:
                    source_seen.add(node)
                    source_nodes.append(node)
        else:
            for parent in parents_to_walk:
                if parent not in completed:
                    queue.append((parent, new_active_path))

    # Fix 4: Ensure every lookup instance in the mapping appears in lookup_used
    # even when the backward traversal did not reach it (e.g. because it feeds
    # into a custom Union Transformation that was not recognised as a branching
    # node before Fix 3, or because it is an unconnected lookup referenced only
    # in an expression not yet parsed).
    already_in_lookup_used: set[str] = {str(lk.get("instance_name", "")) for lk in lookup_used}
    # Build a fast forward-adjacency map: instance -> set of downstream instances
    _fwd_adj: dict[str, set[str]] = defaultdict(set)
    for (f_inst, _), to_list in idx.forward_edges.items():
        for to_inst, _ in to_list:
            _fwd_adj[f_inst].add(to_inst)
    target_of_start = start_node[0]
    for lk_meta in idx.all_lookup_instances():
        lk_inst = lk_meta["instance_name"]
        if lk_inst in already_in_lookup_used:
            continue
        # BFS forward from lk_inst to check whether target_of_start is reachable
        fwd_visited: set[str] = set()
        bfs_fwd: deque[str] = deque([lk_inst])
        reachable_fwd = False
        while bfs_fwd and not reachable_fwd:
            cur_inst = bfs_fwd.popleft()
            if cur_inst in fwd_visited:
                continue
            fwd_visited.add(cur_inst)
            if cur_inst == target_of_start:
                reachable_fwd = True
                break
            for nxt in _fwd_adj.get(cur_inst, set()):
                if nxt not in fwd_visited:
                    bfs_fwd.append(nxt)
        if reachable_fwd:
            already_in_lookup_used.add(lk_inst)
            lookup_used.append(lk_meta)

    result = {
        "source_nodes": list(dict.fromkeys(reversed(source_nodes))),
        "path_nodes": list(reversed(visit_order)),
        "transformations_touched": list(reversed(touched)),
        "lookup_used": lookup_used,
        "filter_or_router_conditions": conditions_applied,
        "stop_expressions": stop_expressions,
        "unresolved_nodes": unresolved_nodes,
    }

    # Cache result (Strategy §4)
    if lineage_cache is not None:
        lineage_cache.put(start_node, result)

    return result


# ============================================================
# COMPLETENESS VALIDATION  (Strategy §7)
# ============================================================

def _validate_lineage_completeness(
    idx: _MappingIndex,
    lineage_records: list[dict[str, Any]],
    target_instances: list[tuple[str, dict[str, str]]],
    diagnostics: _Diagnostics,
) -> dict[str, Any]:
    """
    Before returning results: (Strategy §7)
    - Verify every target field has lineage.
    - Verify all reachable upstream connectors were visited.
    - Flag unresolved nodes instead of silently skipping them.

    IMPORTANT: "expected fields" is sourced from the TARGET definition's own
    field list (idx.target_fields_by_definition), not merely from connector
    edges. A target field with zero inbound connectors is the most severe
    form of lineage loss — it must never be invisible to validation simply
    because no edge exists to derive it from.
    """
    target_instance_names = [ti for ti, _ in target_instances]
    reachable_nodes = idx.all_reachable_connectors_from_targets(target_instance_names)

    # All target fields that should have lineage records — sourced from the
    # authoritative TARGET definition, falling back to connector-derived
    # fields only if the definition itself could not be resolved.
    expected_fields: set[tuple[str, str]] = set()
    for ti, ti_attrs in target_instances:
        def_name = ti_attrs.get("TRANSFORMATION_NAME") or ti
        def_fields = idx.target_fields_by_definition.get(def_name)
        if def_fields is None:
            def_fields = idx.target_fields_by_definition.get(ti)
        if def_fields:
            for field in def_fields:
                expected_fields.add((ti, field))
        else:
            # No TARGET definition found at all — fall back to whatever
            # connectors exist so we don't silently report zero expected.
            for field in idx.incoming_target_fields.get(ti, set()):
                expected_fields.add((ti, field))

    # Collected lineage records
    actual_fields: set[tuple[str, str]] = {
        (r["target_instance"], r["target_field"]) for r in lineage_records
    }

    missing_lineage: list[tuple[str, str]] = []
    for ef in expected_fields:
        if ef not in actual_fields:
            ti, field = ef
            has_connector = field in idx.incoming_target_fields.get(ti, set())
            diagnostics.warn(
                "missing_connector",
                idx.mapping_name,
                ti,
                field,
                "Target field has no lineage record — no inbound connector found"
                if not has_connector
                else "Target field has no lineage record — may be missing connector",
            )
            missing_lineage.append(ef)

    # Verify visited nodes cover all reachable connector nodes
    visited_nodes: set[tuple[str, str]] = set()
    for rec in lineage_records:
        for path_node in rec.get("path_nodes", []):
            if isinstance(path_node, (list, tuple)) and len(path_node) == 2:
                visited_nodes.add(tuple(path_node))  # type: ignore

    unvisited_reachable = reachable_nodes - visited_nodes
    if unvisited_reachable:
        for inst_name, field_name in sorted(unvisited_reachable)[:20]:  # cap diagnostic noise
            diagnostics.warn(
                "missing_connector",
                idx.mapping_name,
                inst_name,
                field_name,
                "Reachable connector node was not visited during traversal",
            )

    return {
        "expected_target_fields": len(expected_fields),
        "actual_lineage_records": len(actual_fields),
        "missing_lineage_fields": missing_lineage,
        "reachable_nodes_count": len(reachable_nodes),
        "visited_nodes_count": len(visited_nodes),
        "unvisited_reachable_count": len(unvisited_reachable),
    }


def _build_completeness_report(
    mapping_name: str,
    folder_name: str,
    validation: dict[str, Any],
    mapping_lineage_records: list[dict[str, Any]],
    diagnostics: _Diagnostics,
) -> dict[str, Any]:
    """
    Final safeguard (per addendum): instead of silently returning partial
    lineage, build a structured report listing every unresolved node,
    missing connector, unsupported construct, and incomplete path for a
    mapping. Used to decide whether the mapping passes strict completeness
    and, if not, to give the caller everything needed to retry or review.
    """
    unresolved_by_field: list[dict[str, Any]] = []
    for rec in mapping_lineage_records:
        unresolved = rec.get("unresolved_nodes", [])
        if unresolved:
            unresolved_by_field.append(
                {
                    "target_instance": rec["target_instance"],
                    "target_field": rec["target_field"],
                    "unresolved_nodes": unresolved,
                }
            )

    mapping_diag_entries = [
        e for e in diagnostics.all_entries() if e["mapping"] == mapping_name
    ]

    is_complete = (
        not validation["missing_lineage_fields"]
        and validation["unvisited_reachable_count"] == 0
        and not unresolved_by_field
    )

    return {
        "folder": folder_name,
        "mapping": mapping_name,
        "is_complete": is_complete,
        "expected_target_fields": validation["expected_target_fields"],
        "actual_lineage_records": validation["actual_lineage_records"],
        "missing_lineage_fields": [
            {"instance": ti, "field": field}
            for ti, field in validation["missing_lineage_fields"]
        ],
        "reachable_nodes_count": validation["reachable_nodes_count"],
        "visited_nodes_count": validation["visited_nodes_count"],
        "unvisited_reachable_count": validation["unvisited_reachable_count"],
        "fields_with_unresolved_nodes": unresolved_by_field,
        "diagnostic_entries": mapping_diag_entries,
    }


# ============================================================
# SUPPLEMENTARY LINEAGE BUILDERS (retained from v1 with diagnostics)
# ============================================================

def _condition_query_field_derivations(
    idx: _MappingIndex,
    transformations_touched: list[dict[str, Any]],
    lineage_cache: _LineageCache | None = None,
    diagnostics: _Diagnostics | None = None,
) -> list[dict[str, Any]]:
    derivations: list[dict[str, Any]] = []
    seen: set[tuple[str, str, str]] = set()

    for tf in transformations_touched:
        inst_name = str(tf.get("instance_name", ""))
        if not inst_name:
            continue
        inst_attrs = idx.instances.get(inst_name, {})
        if not _is_transformation_instance(inst_attrs):
            continue

        tx_name = idx.transformation_name_for_instance(inst_name)
        tx_fields = idx.transform_field_by_name.get(tx_name, {})
        if not tx_fields:
            continue
        by_upper = {name.upper(): name for name in tx_fields}

        for attr_name, attr_value in idx.transform_table_attributes.get(tx_name, {}).items():
            if not _is_condition_or_query_attribute(attr_name):
                continue
            value = str(attr_value or "")
            if not value.strip():
                continue

            field_names: list[str] = []
            field_seen: set[str] = set()
            for token in re.findall(r"[A-Za-z_][A-Za-z0-9_$#.]*", value):
                name = token.split(".")[-1]
                canon = by_upper.get(name.upper())
                if canon and canon.upper() not in field_seen:
                    field_seen.add(canon.upper())
                    field_names.append(canon)

            for field_name in field_names:
                key = (inst_name.upper(), attr_name.upper(), field_name.upper())
                if key in seen:
                    continue
                seen.add(key)

                trace = _trace_lineage_from_node(
                    idx, (inst_name, field_name), lineage_cache, diagnostics
                )
                field_rule = trace["stop_expressions"][0]["expression"] if trace["stop_expressions"] else ""
                entry: dict[str, Any] = {
                    "instance_name": inst_name,
                    "transform_name": tx_name,
                    "transform_type": _tx_type(inst_attrs),
                    "attribute_name": attr_name,
                    "attribute_value": value,
                    "field_name": field_name,
                    "field_porttype": tx_fields.get(field_name, {}).get("PORTTYPE", ""),
                    "field_ultimate_source_fields": [
                        f"{i}.{f}" for i, f in trace["source_nodes"] if i and f
                    ],
                    "field_lineage_path": [
                        f"{i}.{f}" for i, f in trace["path_nodes"] if i and f
                    ],
                    "field_rule_derivation": field_rule or "Direct pass-through",
                    "field_transformations_touched": [
                        f"{t.get('instance_name', '')}.{t.get('field_name', '')} [{t.get('transform_type', '')}]".strip()
                        for t in trace["transformations_touched"]
                    ],
                    "lookup_used": trace.get("lookup_used", []),
                    "filter_or_router_conditions": trace.get("filter_or_router_conditions", []),
                }

                if not entry["field_ultimate_source_fields"] and any(
                    k in attr_name.lower() for k in ("sql", "query", "override")
                ):
                    sql_hint = _sql_alias_derivation(value, field_name)
                    if sql_hint:
                        entry["field_sql_derivation"] = [sql_hint]

                derivations.append(entry)

    return derivations


def _involved_transformation_field_seeds(
    idx: _MappingIndex,
    transformations_touched: list[dict[str, Any]],
    lookup_condition_derivations: list[dict[str, Any]],
    condition_query_derivations: list[dict[str, Any]],
) -> list[tuple[str, str]]:
    seed_by_instance: dict[str, set[str]] = defaultdict(set)

    for tf in transformations_touched:
        inst = str(tf.get("instance_name", ""))
        field = str(tf.get("field_name", ""))
        if inst and field:
            seed_by_instance[inst].add(field)

    for drv in lookup_condition_derivations:
        inst = str(drv.get("lookup_instance_name", ""))
        for field in (str(drv.get("lookup_field", "")), str(drv.get("driving_field", ""))):
            if inst and field:
                seed_by_instance[inst].add(field)

    for drv in condition_query_derivations:
        inst = str(drv.get("instance_name", ""))
        field = str(drv.get("field_name", ""))
        if inst and field:
            seed_by_instance[inst].add(field)

    involved: list[tuple[str, str]] = []
    seen: set[tuple[str, str]] = set()

    for inst_name, seeds in seed_by_instance.items():
        inst_attrs = idx.instances.get(inst_name, {})
        if not _is_transformation_instance(inst_attrs):
            continue

        tx_name = idx.transformation_name_for_instance(inst_name)
        tx_fields = idx.transform_field_by_name.get(tx_name, {})
        if not tx_fields:
            continue

        field_names = list(tx_fields.keys())
        by_upper = {n.upper(): n for n in field_names}
        adjacency: dict[str, set[str]] = defaultdict(set)

        for field_name, attrs in tx_fields.items():
            expr = str(attrs.get("EXPRESSION", ""))
            for dep in _expression_port_dependencies(expr, field_names):
                if dep == field_name:
                    continue
                adjacency[field_name].add(dep)
                adjacency[dep].add(field_name)

        tx_type = _tx_type(inst_attrs)
        if _is_lookup_type(tx_type):
            lk_cond = idx.transform_table_attributes.get(tx_name, {}).get("Lookup condition", "")
            for left_raw, right_raw in _lookup_condition_pairs(lk_cond):
                left = by_upper.get(left_raw.upper())
                right = by_upper.get(right_raw.upper())
                if left and right and left != right:
                    adjacency[left].add(right)
                    adjacency[right].add(left)

        for attr_name, attr_value in idx.transform_table_attributes.get(tx_name, {}).items():
            if not _is_condition_or_query_attribute(attr_name):
                continue
            refs: list[str] = []
            refs_seen: set[str] = set()
            for token in re.findall(r"[A-Za-z_][A-Za-z0-9_$#.]*", str(attr_value or "")):
                name = token.split(".")[-1]
                canon = by_upper.get(name.upper())
                if canon and canon.upper() not in refs_seen:
                    refs_seen.add(canon.upper())
                    refs.append(canon)
            for i in range(len(refs)):
                for j in range(i + 1, len(refs)):
                    a, b = refs[i], refs[j]
                    adjacency[a].add(b)
                    adjacency[b].add(a)

        queue = deque([f for f in seeds if f in tx_fields])
        visited_local: set[str] = set()
        while queue:
            cur = queue.popleft()
            if cur in visited_local:
                continue
            visited_local.add(cur)
            key = (inst_name, cur)
            if key not in seen:
                seen.add(key)
                involved.append(key)
            for nxt in adjacency.get(cur, set()):
                if nxt not in visited_local:
                    queue.append(nxt)

    return involved


def _expression_field_lineage_breakdown(
    idx: _MappingIndex,
    stop_expressions: list[dict[str, str]],
    transformations_touched: list[dict[str, Any]],
    lineage_cache: _LineageCache | None = None,
    diagnostics: _Diagnostics | None = None,
) -> list[dict[str, Any]]:
    breakdown: list[dict[str, Any]] = []
    seen_expressions: set[str] = set()

    for expr_record in stop_expressions:
        expr = str(expr_record.get("expression", ""))
        inst_name = str(expr_record.get("instance_name", ""))
        field_name = str(expr_record.get("field_name", ""))

        expr_key = expr.upper()
        if expr_key in seen_expressions or not expr:
            continue
        seen_expressions.add(expr_key)

        field_refs: list[str] = []
        field_refs_seen: set[str] = set()
        for token in re.findall(r"[A-Za-z_][A-Za-z0-9_$#.]*", expr):
            name = token.split(".")[-1]
            if name.upper() not in field_refs_seen:
                field_refs_seen.add(name.upper())
                field_refs.append(name)

        field_lineages: list[dict[str, Any]] = []
        for ref_field in field_refs:
            for tf in transformations_touched:
                tf_inst = str(tf.get("instance_name", ""))
                tf_field = str(tf.get("field_name", ""))
                if tf_field.upper() == ref_field.upper():
                    trace = _trace_lineage_from_node(
                        idx, (tf_inst, tf_field), lineage_cache, diagnostics
                    )
                    trace_rule = trace["stop_expressions"][0]["expression"] if trace["stop_expressions"] else ""
                    field_lineages.append({
                        "field_name": ref_field,
                        "field_instance": tf_inst,
                        "ultimate_source_fields": [f"{i}.{f}" for i, f in trace["source_nodes"] if i and f],
                        "lineage_path": [f"{i}.{f}" for i, f in trace["path_nodes"] if i and f],
                        "transformations_touched": [
                            f"{t.get('instance_name', '')}.{t.get('field_name', '')} [{t.get('transform_type', '')}]".strip()
                            for t in trace["transformations_touched"]
                        ],
                        "rule": trace_rule or "Direct pass-through",
                    })
                    break

        if field_lineages:
            breakdown.append({
                "expression": expr,
                "expression_instance": inst_name,
                "expression_field": field_name,
                "field_references_lineage": field_lineages,
            })

    return breakdown


def _lookup_condition_field_derivations(
    idx: _MappingIndex,
    lookup_instance_name: str,
    lookup_transform_name: str,
    lookup_condition: str,
    lineage_cache: _LineageCache | None = None,
    diagnostics: _Diagnostics | None = None,
) -> list[dict[str, Any]]:
    tx_fields = idx.transform_field_by_name.get(lookup_transform_name, {})
    if not tx_fields:
        return []

    by_upper = {name.upper(): name for name in tx_fields}
    derivations: list[dict[str, Any]] = []

    for left_raw, right_raw in _lookup_condition_pairs(lookup_condition):
        left = by_upper.get(left_raw.upper())
        right = by_upper.get(right_raw.upper())
        if not left or not right:
            if diagnostics:
                diagnostics.warn(
                    "unresolved_expression",
                    idx.mapping_name,
                    lookup_instance_name,
                    left_raw or right_raw,
                    f"Lookup condition pair ({left_raw}, {right_raw}) references unknown port(s)",
                )
            continue

        left_porttype = tx_fields.get(left, {}).get("PORTTYPE", "")
        right_porttype = tx_fields.get(right, {}).get("PORTTYPE", "")
        left_is_lookup = "LOOKUP" in left_porttype.upper()
        right_is_lookup = "LOOKUP" in right_porttype.upper()

        def _build_driving_entry(lookup_field: str, driving_field: str, lk_pt: str, dr_pt: str) -> dict[str, Any]:
            entry: dict[str, Any] = {
                "lookup_instance_name": lookup_instance_name,
                "lookup_transform_name": lookup_transform_name,
                "lookup_field": lookup_field,
                "driving_field": driving_field,
                "lookup_field_porttype": lk_pt,
                "driving_field_porttype": dr_pt,
                "condition": f"{lookup_field} = {driving_field}",
            }
            if lookup_instance_name:
                trace = _trace_lineage_from_node(
                    idx, (lookup_instance_name, driving_field), lineage_cache, diagnostics
                )
                driving_rule = trace["stop_expressions"][0]["expression"] if trace["stop_expressions"] else ""
                entry["driving_field_ultimate_source_fields"] = [
                    f"{i}.{f}" for i, f in trace["source_nodes"] if i and f
                ]
                entry["driving_field_lineage_path"] = [
                    f"{i}.{f}" for i, f in trace["path_nodes"] if i and f
                ]
                entry["driving_field_rule_derivation"] = driving_rule or "Direct pass-through"
                entry["driving_field_transformations_touched"] = [
                    f"{t.get('instance_name', '')}.{t.get('field_name', '')} [{t.get('transform_type', '')}]".strip()
                    for t in trace["transformations_touched"]
                ]
                if not entry["driving_field_ultimate_source_fields"]:
                    sql_hints = []
                    for lk in trace.get("lookup_used", []):
                        hint = _sql_alias_derivation(lk.get("lookup_sql_override", ""), driving_field)
                        if hint:
                            sql_hints.append(hint)
                    if sql_hints:
                        entry["driving_field_lookup_sql_derivation"] = sql_hints
            return entry

        if left_is_lookup and not right_is_lookup:
            derivations.append(_build_driving_entry(left, right, left_porttype, right_porttype))
        elif right_is_lookup and not left_is_lookup:
            derivations.append(_build_driving_entry(right, left, right_porttype, left_porttype))
        else:
            derivations.append({
                "lookup_instance_name": lookup_instance_name,
                "lookup_transform_name": lookup_transform_name,
                "lookup_field": left,
                "driving_field": right,
                "lookup_field_porttype": left_porttype,
                "driving_field_porttype": right_porttype,
                "condition": f"{left} = {right}",
            })

    return derivations


# ============================================================
# MULTI-PHASE GRAPH BUILDING  (Strategy §1, §2)
# ============================================================

def _build_port_dependency_graph(
    idx: _MappingIndex,
) -> dict[str, dict[str, list[str]]]:
    """
    Phase 3: Build port-level dependency graph for all transformations.
    Returns {tx_name: {field_name: [dep_field_name, ...]}}
    (Strategy §1 Phase 3)
    """
    graph: dict[str, dict[str, list[str]]] = {}
    for tx_name, tx_fields in idx.transform_field_by_name.items():
        field_names = list(tx_fields.keys())
        port_deps: dict[str, list[str]] = {}
        for field_name, attrs in tx_fields.items():
            expr = str(attrs.get("EXPRESSION", ""))
            deps = _expression_port_dependencies(expr, field_names)
            port_deps[field_name] = [d for d in deps if d != field_name]
        graph[tx_name] = port_deps
    return graph


def _build_expression_dependency_map(
    idx: _MappingIndex,
    port_dependency_graph: dict[str, dict[str, list[str]]],
) -> dict[str, dict[str, str]]:
    """
    Phase 4: Full expression dependency map (tx_name -> {field -> expression}).
    (Strategy §1 Phase 4)
    """
    expr_map: dict[str, dict[str, str]] = {}
    for tx_name, tx_fields in idx.transform_field_by_name.items():
        expr_map[tx_name] = {
            field: str(attrs.get("EXPRESSION", ""))
            for field, attrs in tx_fields.items()
        }
    return expr_map


# ============================================================
# PRIMARY EXTRACTION FUNCTION  (Strategy §1)
# ============================================================

def extract_target_bound_fields(
    mapping_json: dict[str, Any],
    *,
    final_target_definitions: list[str] | tuple[str, ...] | None = None,
    include_all_targets: bool = False,
    checkpoint_dir: str | Path | None = None,
    enable_checkpoints: bool = True,
    enable_cache: bool = True,
    strict_completeness: bool = False,
    completeness_report_path: str | Path | None = None,
) -> list[dict[str, Any]]:
    """
    Multi-phase, zero-loss lineage extraction from Informatica canonical JSON.

    Phases (Strategy §1):
      1. Parse canonical JSON
      2. Build intermediate graph representation (connector + port dependency)
      3. Resolve internal transformation dependencies
      4. Resolve connector dependencies
      5. Compute lineage (with memoization, path-aware visited tracking)
      6. Validate completeness before emitting results

    Additional reliability:
      - Persisted checkpoints per mapping (Strategy §2, §3)
      - Graph memoization cache (Strategy §4)
      - Robust traversal for opaque types (Strategy §5)
      - Path-aware cycle detection (Strategy §6)
      - Completeness validation (Strategy §7)
      - Diagnostics for all unresolved constructs (Strategy §8)

    Final safeguard (addendum):
      - strict_completeness: if True, do NOT silently return partial lineage.
        Any mapping with missing target-field lineage, unvisited reachable
        connectors, or unresolved nodes causes extraction to raise
        `IncompleteLineageError` instead of returning results. The error
        carries a structured `report` (also written to
        `completeness_report_path` if given, and always to
        `_diagnostics_collector` for retrieval) listing every unresolved
        node, missing connector, and incomplete path per mapping so the
        run can be retried or reviewed rather than trusted as best-effort.
      - completeness_report_path: if set, the full per-mapping completeness
        report is written here as JSON regardless of pass/fail, so callers
        always have an audit trail even when strict_completeness is False.
    """
    _log_info("=" * 70)
    _log_info("Starting lineage extraction — Zero-Loss Strategy v2")
    _log_info(f"  Final target definitions: {final_target_definitions}")
    _log_info(f"  Include all targets: {include_all_targets}")
    _log_info(f"  Checkpoints enabled: {enable_checkpoints}")
    _log_info(f"  Cache enabled: {enable_cache}")
    _log_info("=" * 70)

    # ---- Phase 1: Parse root ---
    root = _root_node(mapping_json)
    parent_map = _build_parent_map(root)
    terminal_workflow_mappings = _workflow_terminal_session_mappings(root)
    _log_success(f"Phase 1 complete — {len(terminal_workflow_mappings)} terminal mappings identified")

    diagnostics = _Diagnostics()
    checkpoint_mgr = _CheckpointManager(checkpoint_dir) if enable_checkpoints else None
    global_cache = _LineageCache() if enable_cache else None

    lineage_records: list[dict[str, Any]] = []
    mapping_count = 0
    skipped_count = 0
    completeness_reports: list[dict[str, Any]] = []

    mapping_nodes = [n for n in _iter_nodes(root) if n.get("tag") == "MAPPING"]
    _log_info(f"Found {len(mapping_nodes)} total mappings")
    _log_info("-" * 70)

    for mapping_idx_num, mapping_node in enumerate(mapping_nodes, 1):
        folder_node = _ancestor(mapping_node, parent_map, "FOLDER") or root
        idx = _MappingIndex(mapping_node, folder_node)
        _log_progress(mapping_idx_num, len(mapping_nodes), f"Processing: {idx.mapping_name}")
        _log_info(f"  Folder: {idx.folder_name}")

        # Skip non-terminal mappings (unless explicitly scoped)
        if (
            not include_all_targets
            and not final_target_definitions
            and terminal_workflow_mappings
            and idx.mapping_name not in terminal_workflow_mappings
        ):
            _log_warning(f"  Skipped (not in terminal session mappings)")
            skipped_count += 1
            continue

        mapping_count += 1

        # ---- Phase 2: Build intermediate graph (checkpoint) ----
        if checkpoint_mgr:
            connector_graph_data = {
                "backward_edges": {
                    f"{k[0]}/{k[1]}": [[p[0], p[1]] for p in v]
                    for k, v in idx.backward_edges.items()
                },
                "total_connectors": len(idx.all_connectors),
                "total_instances": len(idx.instances),
            }
            checkpoint_mgr.save("connector_graph", idx.mapping_name, connector_graph_data)

        # ---- Phase 3: Port dependency graph ----
        port_dep_graph = _build_port_dependency_graph(idx)
        if checkpoint_mgr:
            checkpoint_mgr.save("port_dependency_graph", idx.mapping_name, port_dep_graph)
            checkpoint_mgr.validate_counts(
                "port_dependency_graph",
                idx.mapping_name,
                len(idx.transformations),
                len(port_dep_graph),
            )

        # ---- Phase 4: Expression dependency map ----
        expr_dep_map = _build_expression_dependency_map(idx, port_dep_graph)
        if checkpoint_mgr:
            checkpoint_mgr.save("expression_dep_map", idx.mapping_name, expr_dep_map)

        # ---- Lookup / Router / Joiner metadata checkpoints ----
        if checkpoint_mgr:
            lookup_meta = {}
            router_meta = {}
            joiner_meta = {}
            for tx_name, tx_node in idx.transformations.items():
                tx_type = _attr(tx_node, "TYPE", "").upper()
                ta = idx.transform_table_attributes.get(tx_name, {})
                if "LOOKUP" in tx_type:
                    lookup_meta[tx_name] = {
                        "table": ta.get("Lookup table name", ""),
                        "condition": ta.get("Lookup condition", ""),
                    }
                elif "ROUTER" in tx_type:
                    router_meta[tx_name] = {k: v for k, v in ta.items() if "condition" in k.lower()}
                elif "JOINER" in tx_type:
                    joiner_meta[tx_name] = {
                        "join_condition": ta.get("Join Condition", ""),
                        "join_type": ta.get("Join Type", ""),
                    }
            checkpoint_mgr.save("lookup_metadata", idx.mapping_name, lookup_meta)
            checkpoint_mgr.save("router_group_metadata", idx.mapping_name, router_meta)
            checkpoint_mgr.save("joiner_metadata", idx.mapping_name, joiner_meta)

        # ---- Resolve target instances ----
        scoped_defs_upper = {
            str(name).strip().upper()
            for name in (final_target_definitions or [])
            if str(name).strip()
        }

        target_instances: list[tuple[str, dict[str, str]]] = []
        if include_all_targets:
            for inst_name, inst_attrs in idx.instances.items():
                if _is_target_instance(inst_attrs):
                    target_instances.append((inst_name, inst_attrs))
        elif scoped_defs_upper:
            for inst_name, inst_attrs in idx.instances.items():
                if not _is_target_instance(inst_attrs):
                    continue
                def_name = inst_attrs.get("TRANSFORMATION_NAME") or inst_name
                if def_name.upper() in scoped_defs_upper:
                    target_instances.append((inst_name, inst_attrs))
            if not target_instances:
                continue
        else:
            for inst_name, inst_attrs in idx.instances.items():
                if not _is_target_instance(inst_attrs):
                    continue
                def_name = inst_attrs.get("TRANSFORMATION_NAME") or inst_name
                if def_name.upper() in idx.final_target_definition_names_upper:
                    target_instances.append((inst_name, inst_attrs))

            if not target_instances and idx.target_load_order:
                max_order = max(order for order, _ in idx.target_load_order)
                ordered_targets = [ti for order, ti in idx.target_load_order if order == max_order]
                seen_ti: set[str] = set()
                for ti in ordered_targets:
                    if ti in seen_ti:
                        continue
                    seen_ti.add(ti)
                    inst_attrs = idx.instances.get(ti, {})
                    if _is_target_instance(inst_attrs):
                        target_instances.append((ti, inst_attrs))

            if not target_instances:
                for inst_name, inst_attrs in idx.instances.items():
                    if _is_target_instance(inst_attrs) and idx.incoming_target_fields.get(inst_name):
                        target_instances.append((inst_name, inst_attrs))
                if not target_instances:
                    for inst_name, inst_attrs in idx.instances.items():
                        if _is_target_instance(inst_attrs):
                            target_instances.append((inst_name, inst_attrs))

        # ---- Phase 5: Compute lineage per target field ----
        mapping_lineage_records: list[dict[str, Any]] = []
        mapping_cache = _LineageCache() if enable_cache else None

        for target_instance, target_inst_attrs in target_instances:
            def_name = target_inst_attrs.get("TRANSFORMATION_NAME") or target_instance
            fields_from_definition = idx.target_fields_by_definition.get(def_name, [])
            all_fields = fields_from_definition or sorted(
                idx.incoming_target_fields.get(target_instance, set())
            )
            _log_info(f"  Target: {target_instance} | Fields: {len(all_fields)}")

            for tf_idx_num, target_field in enumerate(all_fields, 1):
                # Use per-mapping cache for sub-calls; fall back to global
                active_cache = mapping_cache or global_cache

                trace = _trace_lineage_from_node(
                    idx,
                    (target_instance, target_field),
                    lineage_cache=active_cache,
                    diagnostics=diagnostics,
                )

                source_nodes = trace["source_nodes"]
                path_nodes = trace["path_nodes"]
                transformations_touched = trace["transformations_touched"]
                lookup_used = trace["lookup_used"]
                conditions_applied = trace["filter_or_router_conditions"]
                stop_expressions = trace["stop_expressions"]
                unresolved = trace.get("unresolved_nodes", [])

                primary_expression = stop_expressions[0]["expression"] if stop_expressions else ""
                source_field_names = [f"{i}.{f}" for i, f in source_nodes if i and f]

                lookup_condition_derivations: list[dict[str, Any]] = []
                for lk in lookup_used:
                    lookup_condition_derivations.extend(
                        _lookup_condition_field_derivations(
                            idx,
                            lk.get("instance_name", ""),
                            lk.get("transform_name", ""),
                            lk.get("lookup_condition", ""),
                            lineage_cache=active_cache,
                            diagnostics=diagnostics,
                        )
                    )

                condition_query_derivations = _condition_query_field_derivations(
                    idx, transformations_touched, active_cache, diagnostics
                )

                involved_field_seeds = _involved_transformation_field_seeds(
                    idx, transformations_touched, lookup_condition_derivations, condition_query_derivations
                )

                expression_field_lineage = _expression_field_lineage_breakdown(
                    idx, stop_expressions, transformations_touched, active_cache, diagnostics
                )

                # Build transformation field lineage section
                transformation_field_lineage_section: list[dict[str, Any]] = []
                seed_seen: set[tuple[str, str]] = set()
                section_seeds: list[tuple[str, str]] = []

                for tf in transformations_touched:
                    seed = (tf.get("instance_name", ""), tf.get("field_name", ""))
                    if seed[0] and seed[1] and seed not in seed_seen:
                        seed_seen.add(seed)
                        section_seeds.append(seed)

                for drv in lookup_condition_derivations:
                    lk_inst = str(drv.get("lookup_instance_name", ""))
                    lk_field = str(drv.get("lookup_field", ""))
                    dr_field = str(drv.get("driving_field", ""))
                    for seed in ((lk_inst, lk_field), (lk_inst, dr_field)):
                        if seed[0] and seed[1] and seed not in seed_seen:
                            seed_seen.add(seed)
                            section_seeds.append(seed)

                for drv in condition_query_derivations:
                    seed = (str(drv.get("instance_name", "")), str(drv.get("field_name", "")))
                    if seed[0] and seed[1] and seed not in seed_seen:
                        seed_seen.add(seed)
                        section_seeds.append(seed)

                for seed in section_seeds:
                    sub_trace = _trace_lineage_from_node(idx, seed, active_cache, diagnostics)
                    sub_rule = sub_trace["stop_expressions"][0]["expression"] if sub_trace["stop_expressions"] else ""
                    transformation_field_lineage_section.append({
                        "Target Field": f"{seed[0]}.{seed[1]}",
                        "Ultimate Source Field(s)": [f"{i}.{f}" for i, f in sub_trace["source_nodes"] if i and f],
                        "Lineage Path": [f"{i}.{f}" for i, f in sub_trace["path_nodes"] if i and f],
                        "Transformation(s) Touched": [
                            f"{t.get('instance_name', '')}.{t.get('field_name', '')} [{t.get('transform_type', '')}]".strip()
                            for t in sub_trace["transformations_touched"]
                        ],
                        "Rule / Derivation": sub_rule or "Direct pass-through",
                        "Lookup Used": sub_trace["lookup_used"],
                        "Filter / Router Conditions": sub_trace["filter_or_router_conditions"],
                    })

                additional_involved_field_lineage_section: list[dict[str, Any]] = []
                add_seen: set[tuple[str, str]] = set(section_seeds)
                for seed in involved_field_seeds:
                    if seed in add_seen:
                        continue
                    add_seen.add(seed)
                    sub_trace = _trace_lineage_from_node(idx, seed, active_cache, diagnostics)
                    sub_rule = sub_trace["stop_expressions"][0]["expression"] if sub_trace["stop_expressions"] else ""
                    additional_involved_field_lineage_section.append({
                        "Target Field": f"{seed[0]}.{seed[1]}",
                        "Ultimate Source Field(s)": [f"{i}.{f}" for i, f in sub_trace["source_nodes"] if i and f],
                        "Lineage Path": [f"{i}.{f}" for i, f in sub_trace["path_nodes"] if i and f],
                        "Transformation(s) Touched": [
                            f"{t.get('instance_name', '')}.{t.get('field_name', '')} [{t.get('transform_type', '')}]".strip()
                            for t in sub_trace["transformations_touched"]
                        ],
                        "Rule / Derivation": sub_rule or "Direct pass-through",
                        "Lookup Used": sub_trace["lookup_used"],
                        "Filter / Router Conditions": sub_trace["filter_or_router_conditions"],
                    })

                lookup_condition_field_lineage_section: list[dict[str, Any]] = []
                lc_seen: set[tuple[str, str, str]] = set()
                for drv in lookup_condition_derivations:
                    lk_inst = str(drv.get("lookup_instance_name", ""))
                    lk_field = str(drv.get("lookup_field", ""))
                    dr_field = str(drv.get("driving_field", ""))
                    condition = str(drv.get("condition", ""))
                    if not lk_inst or not dr_field:
                        continue
                    key = (lk_inst.upper(), dr_field.upper(), condition.upper())
                    if key in lc_seen:
                        continue
                    lc_seen.add(key)
                    lookup_condition_field_lineage_section.append({
                        "Target Field": f"{lk_inst}.{dr_field}",
                        "Lookup Condition": condition,
                        "Lookup Field": f"{lk_inst}.{lk_field}" if lk_field else "",
                        "Ultimate Source Field(s)": list(drv.get("driving_field_ultimate_source_fields", [])),
                        "Lineage Path": list(drv.get("driving_field_lineage_path", [])),
                        "Transformation(s) Touched": list(drv.get("driving_field_transformations_touched", [])),
                        "Rule / Derivation": str(drv.get("driving_field_rule_derivation", "")) or "Direct pass-through",
                        "Lookup Used": [
                            lk for lk in lookup_used
                            if str(lk.get("instance_name", "")).upper() == lk_inst.upper()
                        ],
                        "Filter / Router Conditions": conditions_applied,
                    })

                condition_query_field_lineage_section: list[dict[str, Any]] = []
                cq_seen: set[tuple[str, str, str]] = set()
                for drv in condition_query_derivations:
                    drv_inst = str(drv.get("instance_name", ""))
                    attr_name = str(drv.get("attribute_name", ""))
                    drv_field = str(drv.get("field_name", ""))
                    key = (drv_inst.upper(), attr_name.upper(), drv_field.upper())
                    if key in cq_seen:
                        continue
                    cq_seen.add(key)
                    condition_query_field_lineage_section.append({
                        "Target Field": f"{drv_inst}.{drv_field}",
                        "Transformation": f"{drv_inst} [{drv.get('transform_type', '')}]",
                        "Attribute Name": attr_name,
                        "Attribute Value": str(drv.get("attribute_value", "")),
                        "Ultimate Source Field(s)": list(drv.get("field_ultimate_source_fields", [])),
                        "Lineage Path": list(drv.get("field_lineage_path", [])),
                        "Transformation(s) Touched": list(drv.get("field_transformations_touched", [])),
                        "Rule / Derivation": str(drv.get("field_rule_derivation", "")) or "Direct pass-through",
                        "Lookup Used": list(drv.get("lookup_used", [])),
                        "Filter / Router Conditions": list(drv.get("filter_or_router_conditions", [])),
                        "SQL Derivation": list(drv.get("field_sql_derivation", [])),
                    })

                report_section = {
                    "Target Field": target_field,
                    "Ultimate Source Field(s)": source_field_names,
                    "Lineage Path": [f"{i}.{f}" for i, f in path_nodes if i and f],
                    "Transformation(s) Touched": [
                        f"{t.get('instance_name', '')}.{t.get('field_name', '')} [{t.get('transform_type', '')}]".strip()
                        for t in transformations_touched
                    ],
                    "Rule / Derivation": primary_expression or "Direct pass-through",
                    "Lookup Used": lookup_used,
                    "Filter / Router Conditions": conditions_applied,
                    "Lookup Condition Field Derivations": lookup_condition_derivations,
                    "Condition/Query Field Derivations": condition_query_derivations,
                }

                rec = {
                    "folder": idx.folder_name,
                    "mapping": idx.mapping_name,
                    "target_instance": target_instance,
                    "target_field": target_field,
                    "target_column_name": target_field,
                    "source_nodes": source_nodes,
                    "ultimate_source_fields": source_field_names,
                    "transformations_touched": transformations_touched,
                    "path_nodes": path_nodes,
                    "stop_expressions": stop_expressions,
                    "transformation_rule": primary_expression,
                    "transformation_rule_plain_english": (
                        humanise_rule(primary_expression) if primary_expression else "Direct pass-through"
                    ),
                    "lookup_used": lookup_used,
                    "filter_or_router_conditions": conditions_applied,
                    "lineage_report_section": report_section,
                    "transformation_field_lineage_section": transformation_field_lineage_section,
                    "additional_involved_field_lineage_section": additional_involved_field_lineage_section,
                    "lookup_condition_field_lineage_section": lookup_condition_field_lineage_section,
                    "condition_query_field_lineage_section": condition_query_field_lineage_section,
                    "expression_field_lineage_section": expression_field_lineage,
                    "unresolved_nodes": unresolved,
                }
                mapping_lineage_records.append(rec)
                _log_info(
                    f"    └─ [{tf_idx_num}/{len(all_fields)}] {target_field} "
                    f"({len(source_field_names)} sources, {len(unresolved)} unresolved)"
                )

        # ---- Phase 6: Completeness validation (Strategy §7) ----
        validation = _validate_lineage_completeness(
            idx, mapping_lineage_records, target_instances, diagnostics
        )
        _log_info(
            f"  Validation: expected={validation['expected_target_fields']} "
            f"actual={validation['actual_lineage_records']} "
            f"missing={len(validation['missing_lineage_fields'])} "
            f"unvisited_reachable={validation['unvisited_reachable_count']}"
        )

        mapping_report = _build_completeness_report(
            idx.mapping_name, idx.folder_name, validation, mapping_lineage_records, diagnostics
        )
        completeness_reports.append(mapping_report)
        if not mapping_report["is_complete"]:
            _log_warning(
                f"  Completeness FAILED for mapping '{idx.mapping_name}': "
                f"{len(mapping_report['missing_lineage_fields'])} missing field(s), "
                f"{mapping_report['unvisited_reachable_count']} unvisited reachable node(s), "
                f"{len(mapping_report['fields_with_unresolved_nodes'])} field(s) with unresolved nodes"
            )

        # Save partial lineage cache checkpoint (Strategy §2)
        if checkpoint_mgr and mapping_cache:
            checkpoint_mgr.save(
                "partial_lineage_cache",
                idx.mapping_name,
                {
                    "record_count": len(mapping_lineage_records),
                    "cache_stats": mapping_cache.stats(),
                    "validation": {
                        k: (list(v) if isinstance(v, list) else v)
                        for k, v in validation.items()
                    },
                },
            )

        lineage_records.extend(mapping_lineage_records)

    # Cleanup temp checkpoints (only if we created them internally)
    if checkpoint_mgr and checkpoint_dir is None:
        checkpoint_mgr.cleanup()

    # Final summary
    diag_summary = diagnostics.summary()
    _log_info("-" * 70)
    _log_success("Lineage extraction complete")
    _log_info(f"  Mappings processed: {mapping_count}")
    _log_info(f"  Mappings skipped: {skipped_count}")
    _log_info(f"  Total lineage records: {len(lineage_records)}")
    if diag_summary:
        _log_warning(f"  Diagnostic warnings: {diag_summary}")
    if enable_cache and global_cache:
        _log_info(f"  Cache stats: {global_cache.stats()}")
    _log_info("-" * 70)

    # ---- Final safeguard: completeness report + strict-mode gate (addendum) ----
    incomplete_mappings = [r for r in completeness_reports if not r["is_complete"]]
    full_report = {
        "generated_at": datetime.now().isoformat(),
        "strict_completeness": strict_completeness,
        "total_mappings": len(completeness_reports),
        "incomplete_mapping_count": len(incomplete_mappings),
        "mappings": completeness_reports,
    }

    if completeness_report_path:
        try:
            out_report = Path(completeness_report_path)
            out_report.parent.mkdir(parents=True, exist_ok=True)
            out_report.write_text(
                json.dumps(full_report, indent=2, ensure_ascii=False, default=str),
                encoding="utf-8",
            )
            _log_success(f"Completeness report written: {out_report.absolute()}")
        except Exception as e:
            _log_warning(f"Failed to write completeness report: {e}")

    if incomplete_mappings:
        names = ", ".join(r["mapping"] for r in incomplete_mappings)
        if strict_completeness:
            _log_error(
                f"STRICT COMPLETENESS FAILED for {len(incomplete_mappings)} mapping(s): {names}. "
                "Refusing to return partial lineage — see report for unresolved nodes, "
                "missing connectors, and incomplete paths."
            )
            raise IncompleteLineageError(
                f"Lineage extraction incomplete for {len(incomplete_mappings)} mapping(s): {names}. "
                "Set strict_completeness=False to allow best-effort partial results, "
                "or fix/retry the affected mapping(s). Full diagnostic report attached.",
                full_report,
            )
        else:
            _log_warning(
                f"  {len(incomplete_mappings)} mapping(s) incomplete (best-effort mode): {names}"
            )

    return lineage_records


# ============================================================
# HUMANISE RULE
# ============================================================

_HUMANISE_RULES: list[tuple[re.Pattern[str], str]] = [
    (re.compile(r"^\s*(\w+)\s*$"), r"Pass through \1 without change"),
    (
        re.compile(r"IIF\s*\(\s*(.+?)\s*,\s*(.+?)\s*,\s*(.+?)\s*\)\s*$", re.IGNORECASE),
        r"If \1 then \2, otherwise \3",
    ),
    (re.compile(r"DECODE\s*\(", re.IGNORECASE), "Map value using DECODE logic"),
    (re.compile(r"NVL2?\s*\(", re.IGNORECASE), "Replace NULL values using NVL logic"),
    (re.compile(r"(L?R?TRIM)\s*\(\s*(\w+)\s*\)\s*$", re.IGNORECASE), r"Trim whitespace from \2"),
    (re.compile(r"(UPPER|LOWER)\s*\(\s*(\w+)\s*\)\s*$", re.IGNORECASE), r"Convert \2 to \1 case"),
    (
        re.compile(r"SUBSTR(?:ING)?\s*\(\s*(\w+)\s*,\s*(\d+)\s*(?:,\s*(\d+))?\s*\)", re.IGNORECASE),
        r"Extract part of \1 starting at position \2",
    ),
    (re.compile(r"TO_DATE\s*\(\s*(\w+)\s*,\s*'(.+?)'\s*\)", re.IGNORECASE), r"Parse \1 as date using format '\2'"),
    (re.compile(r"TO_CHAR\s*\(\s*(\w+)", re.IGNORECASE), r"Convert \1 to text"),
    (re.compile(r"\|\|"), "Concatenate values"),
    (re.compile(r"[\+\-\*/]"), "Apply arithmetic calculation"),
    (re.compile(r"\bLKP\.", re.IGNORECASE), "Perform a lookup"),
    (re.compile(r"\bSEQGEN\b|\bNEXTVAL\b", re.IGNORECASE), "Generate next sequence value"),
    (re.compile(r"\bSYSDATE\b|\bSYSTIMESTAMP\b", re.IGNORECASE), "Use current system date/time"),
]


def humanise_rule(expression: str) -> str:
    """Convert an Informatica expression string into plain English."""
    if not expression or not expression.strip():
        return "No transformation logic provided"
    expr = expression.strip()
    for pattern, template in _HUMANISE_RULES:
        m = pattern.search(expr)
        if m:
            try:
                text = m.expand(template)
            except re.error:
                text = template
            if text:
                return text[0].upper() + text[1:]
            return "Apply transformation expression"
    return f"Apply expression: {textwrap.shorten(expr, width=140, placeholder=' ...')}"


# ============================================================
# DOCX BRIEFING BUILDER
# ============================================================

def build_briefing_docx(
    mapping_json: dict[str, Any],
    out_path: str | Path = "fsd_briefing.docx",
    *,
    folder_name: str | None = None,
    mapping_name: str | None = None,
    final_target_definitions: list[str] | tuple[str, ...] | None = None,
    include_all_targets: bool = False,
    checkpoint_dir: str | Path | None = None,
    enable_checkpoints: bool = True,
    enable_cache: bool = True,
    strict_completeness: bool = False,
    completeness_report_path: str | Path | None = None,
) -> Path:
    """
    Build a Word briefing document scoped by optional folder/mapping filters.

    If strict_completeness=True, the underlying extraction will raise
    `IncompleteLineageError` (instead of building a briefing from partial
    lineage) when any in-scope mapping has unresolved nodes, missing
    connectors, or unvisited reachable nodes. The error's `.report`
    attribute contains the full structured diagnostic breakdown.
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required to build .docx output. Install with: pip install python-docx"
        ) from exc

    out = Path(out_path)
    lineage = extract_target_bound_fields(
        mapping_json,
        final_target_definitions=final_target_definitions,
        include_all_targets=include_all_targets,
        checkpoint_dir=checkpoint_dir,
        enable_checkpoints=enable_checkpoints,
        enable_cache=enable_cache,
        strict_completeness=strict_completeness,
        completeness_report_path=completeness_report_path,
    )

    if folder_name:
        lineage = [r for r in lineage if r["folder"] == folder_name]
    if mapping_name:
        lineage = [r for r in lineage if r["mapping"] == mapping_name]

    if not lineage:
        raise ValueError(
            "No lineage records found for the given scope. "
            f"folder_name={folder_name!r}, mapping_name={mapping_name!r}"
        )

    humanised: dict[str, str] = {}
    for rec in lineage:
        for tf in rec["transformations_touched"]:
            key = f"{tf['instance_name']}::{tf['field_name']}"
            if key not in humanised:
                humanised[key] = humanise_rule(tf.get("expression", ""))

    by_mapping: dict[tuple[str, str], list[dict[str, Any]]] = defaultdict(list)
    for rec in lineage:
        by_mapping[(rec["folder"], rec["mapping"])].append(rec)

    doc = Document()
    doc.add_heading("Informatica FSD Briefing", level=1)
    doc.add_paragraph("Generated from canonical Informatica XML JSON — Zero-Loss Strategy v2.")

    for (folder, mapping), records in sorted(by_mapping.items()):
        doc.add_heading(f"Folder: {folder}", level=2)
        doc.add_heading(f"Mapping: {mapping}", level=3)

        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = "Target Field"
        hdr[1].text = "Source Field(s)"
        hdr[2].text = "Transformation Chain"
        hdr[3].text = "Rule (Plain English)"

        for rec in records:
            row = table.add_row().cells
            row[0].text = rec["target_field"]
            row[1].text = "\n".join(rec.get("ultimate_source_fields", [])) or "-"
            row[2].text = " -> ".join(f"{i}.{f}" for i, f in rec["path_nodes"]) or "-"
            row[3].text = rec.get("transformation_rule_plain_english") or "Direct pass-through"

        doc.add_paragraph("")
        doc.add_heading("Transformation Detail", level=4)

        detail = doc.add_table(rows=1, cols=4)
        d = detail.rows[0].cells
        d[0].text = "Instance"
        d[1].text = "Field"
        d[2].text = "Type"
        d[3].text = "Expression / Rule"

        seen: set[tuple[str, str]] = set()
        for rec in records:
            for tf in rec["transformations_touched"]:
                k = (tf["instance_name"], tf["field_name"])
                if k in seen:
                    continue
                seen.add(k)
                dr = detail.add_row().cells
                dr[0].text = tf["instance_name"]
                dr[1].text = tf["field_name"]
                dr[2].text = tf.get("transform_type", "")
                hkey = f"{tf['instance_name']}::{tf['field_name']}"
                dr[3].text = humanised.get(hkey) or tf.get("expression", "") or "-"

        # Additional involved field lineage
        additional_rows: list[dict[str, Any]] = []
        additional_seen: set[tuple[str, str]] = set()
        for rec in records:
            for row in rec.get("additional_involved_field_lineage_section", []):
                target_field_str = str(row.get("Target Field", ""))
                if "." in target_field_str:
                    inst, fld = target_field_str.split(".", 1)
                else:
                    inst, fld = target_field_str, ""
                key = (inst, fld)
                if key in additional_seen:
                    continue
                additional_seen.add(key)
                additional_rows.append(row)

        if additional_rows:
            doc.add_paragraph("")
            doc.add_heading("Additional Involved Field Lineage", level=4)
            extra = doc.add_table(rows=1, cols=4)
            e = extra.rows[0].cells
            e[0].text = "Transformation Field"
            e[1].text = "Ultimate Source Field(s)"
            e[2].text = "Lineage Path"
            e[3].text = "Rule / Derivation"
            for row_data in additional_rows:
                rr = extra.add_row().cells
                rr[0].text = str(row_data.get("Target Field", ""))
                rr[1].text = "\n".join(row_data.get("Ultimate Source Field(s)", [])) or "-"
                rr[2].text = " -> ".join(row_data.get("Lineage Path", [])) or "-"
                rr[3].text = str(row_data.get("Rule / Derivation", "")) or "Direct pass-through"

        # Diagnostics summary in doc
        all_unresolved: list[dict[str, str]] = []
        for rec in records:
            all_unresolved.extend(rec.get("unresolved_nodes", []))
        if all_unresolved:
            doc.add_paragraph("")
            doc.add_heading("Diagnostic Warnings", level=4)
            doc.add_paragraph(
                f"{len(all_unresolved)} unresolved node(s) detected during traversal. "
                "These may indicate missing connectors, opaque transformations, or unsupported constructs."
            )
            warn_table = doc.add_table(rows=1, cols=3)
            wh = warn_table.rows[0].cells
            wh[0].text = "Instance"
            wh[1].text = "Field"
            wh[2].text = "Reason"
            unres_seen: set[tuple[str, str]] = set()
            for u in all_unresolved:
                uk = (u.get("instance", ""), u.get("field", ""))
                if uk in unres_seen:
                    continue
                unres_seen.add(uk)
                wr = warn_table.add_row().cells
                wr[0].text = u.get("instance", "")
                wr[1].text = u.get("field", "")
                wr[2].text = u.get("reason", "")

        doc.add_page_break()

        # Expression Field Lineage Breakdown
        expr_breakdowns: list[dict[str, Any]] = []
        for rec in records:
            expr_breakdowns.extend(rec.get("expression_field_lineage_section", []))

        if expr_breakdowns:
            doc.add_paragraph("")
            doc.add_heading("Expression Field Lineage Breakdown", level=4)
            doc.add_paragraph(
                "Lineage for each field referenced within the target field's transformation expression."
            )
            for expr_bd in expr_breakdowns:
                expr_text = str(expr_bd.get("expression", ""))[:200]
                doc.add_paragraph(f"Expression: {expr_text}", style="List Bullet")
                field_refs = expr_bd.get("field_references_lineage", [])
                if field_refs:
                    expr_table = doc.add_table(rows=1, cols=4)
                    e = expr_table.rows[0].cells
                    e[0].text = "Field Reference"
                    e[1].text = "Ultimate Source Field(s)"
                    e[2].text = "Lineage Path"
                    e[3].text = "Rule / Derivation"
                    for field_ref in field_refs:
                        fr = expr_table.add_row().cells
                        fr[0].text = str(field_ref.get("field_name", ""))
                        fr[1].text = "\n".join(field_ref.get("ultimate_source_fields", [])) or "-"
                        fr[2].text = " -> ".join(field_ref.get("lineage_path", [])[:10]) or "-"
                        fr[3].text = str(field_ref.get("rule", "")) or "Direct pass-through"
                doc.add_paragraph("")

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def build_breifing_docx(*args, **kwargs):
    """Compatibility alias for misspelled function name in requirements."""
    return build_briefing_docx(*args, **kwargs)


# ============================================================
# CLI
# ============================================================

def _cli() -> int:
    start_time = time.time()

    parser = argparse.ArgumentParser(
        description="Build FSD briefing DOCX from canonical Informatica JSON (Zero-Loss Strategy v2)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=textwrap.dedent(
            """
            Examples:
              python fsd_briefing_builder_v2.py -i wf.json -o fsd.docx
              python fsd_briefing_builder_v2.py -i wf.json --folder CBR_FOLDER
              python fsd_briefing_builder_v2.py -i wf.json --mapping m_CBR_DTM_DESTROY_TRADEID
              python fsd_briefing_builder_v2.py -i wf.json --all-mappings --lineage-only --lineage-out-dir output/JSON
              python fsd_briefing_builder_v2.py -i wf.json --lineage-only
              python fsd_briefing_builder_v2.py -i wf.json --checkpoint-dir /tmp/infa_ckpts
              python fsd_briefing_builder_v2.py -i wf.json --no-cache --no-checkpoints
            """
        ),
    )
    parser.add_argument("-i", "--input", required=True, help="Canonical JSON file path")
    parser.add_argument("-o", "--output", default="fsd_briefing.docx", help="Output DOCX path")
    parser.add_argument("--folder", help="Optional folder scope")
    parser.add_argument("--mapping", help="Optional mapping scope")
    parser.add_argument(
        "--all-mappings",
        action="store_true",
        help="Generate one lineage JSON file per mapping.",
    )
    parser.add_argument(
        "--final-target",
        action="append",
        dest="final_targets",
        help=(
            "Final target definition to include. Repeat for multiple values. "
            "If omitted, final target(s) are detected dynamically from TARGETLOADORDER."
        ),
    )
    parser.add_argument(
        "--all-targets",
        action="store_true",
        help="Disable final target filter and include all target definitions.",
    )
    parser.add_argument("--lineage-only", action="store_true", help="Print lineage JSON and exit")
    parser.add_argument("--lineage-out", help="Write lineage JSON to this file path")
    parser.add_argument("--lineage-out-dir", help="Output directory for --all-mappings lineage JSON files.")
    parser.add_argument(
        "--checkpoint-dir",
        default=None,
        help="Persist intermediate checkpoints to this directory (default: temp dir, auto-cleaned).",
    )
    parser.add_argument(
        "--no-checkpoints",
        action="store_true",
        help="Disable intermediate checkpoint persistence.",
    )
    parser.add_argument(
        "--no-cache",
        action="store_true",
        help="Disable lineage graph memoization cache.",
    )
    parser.add_argument(
        "--diagnostics-out",
        default=None,
        help="Write diagnostic warnings JSON to this path.",
    )
    parser.add_argument(
        "--strict-completeness",
        action="store_true",
        help=(
            "Fail the run instead of emitting partial lineage if any target "
            "field cannot be fully resolved. Writes a diagnostic report "
            "listing unresolved nodes, missing connectors, unsupported "
            "expressions, and incomplete paths so the mapping can be "
            "retried or reviewed."
        ),
    )
    parser.add_argument(
        "--completeness-report",
        default=None,
        help=(
            "Path to write the full per-mapping completeness report JSON "
            "(written regardless of pass/fail; required reading when "
            "--strict-completeness causes a failure)."
        ),
    )
    args = parser.parse_args()

    _log_info("#" * 70)
    _log_info("FSD BRIEFING BUILDER — Zero-Loss Strategy v2")
    _log_info("#" * 70)
    _log_info(f"Started at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    _log_info("-" * 70)
    _log_info("ARGUMENTS:")
    _log_info(f"  Input file     : {args.input}")
    _log_info(f"  Output file    : {args.output}")
    _log_info(f"  Folder filter  : {args.folder or 'None'}")
    _log_info(f"  Mapping filter : {args.mapping or 'None'}")
    _log_info(f"  All mappings   : {args.all_mappings}")
    _log_info(f"  Lineage only   : {args.lineage_only}")
    _log_info(f"  Checkpoint dir : {args.checkpoint_dir or '(auto temp)' if not args.no_checkpoints else 'disabled'}")
    _log_info(f"  Cache          : {'disabled' if args.no_cache else 'enabled'}")
    _log_info(f"  Strict complete: {args.strict_completeness}")
    _log_info(f"  Complete report: {args.completeness_report or 'None'}")
    _log_info("-" * 70)

    if args.all_mappings and args.mapping:
        _log_error("--all-mappings cannot be used with --mapping")
        parser.error("--all-mappings cannot be used with --mapping")
    if args.all_mappings and args.lineage_out:
        _log_error("Use --lineage-out-dir (not --lineage-out) with --all-mappings")
        parser.error("Use --lineage-out-dir (not --lineage-out) with --all-mappings")

    _log_info("Loading input JSON file...")
    try:
        data = json.loads(Path(args.input).read_text(encoding="utf-8"))
        _log_success("Successfully loaded JSON file")
    except FileNotFoundError:
        _log_error(f"Input file not found: {args.input}")
        return 1
    except json.JSONDecodeError as e:
        _log_error(f"Invalid JSON format: {e}")
        return 1

    root = _root_node(data)

    enable_checkpoints = not args.no_checkpoints
    enable_cache = not args.no_cache
    checkpoint_dir = args.checkpoint_dir

    if args.all_mappings:
        _log_info("=" * 70)
        _log_info("ALL-MAPPINGS MODE: Generating lineage per mapping")
        _log_info("=" * 70)

        try:
            lineage = extract_target_bound_fields(
                data,
                final_target_definitions=(args.final_targets if not args.all_targets else None),
                include_all_targets=True,
                checkpoint_dir=checkpoint_dir,
                enable_checkpoints=enable_checkpoints,
                enable_cache=enable_cache,
                strict_completeness=args.strict_completeness,
                completeness_report_path=args.completeness_report,
            )
        except IncompleteLineageError as e:
            _log_error(str(e))
            if args.completeness_report:
                _log_error(f"See completeness report for details: {args.completeness_report}")
            else:
                _log_error(
                    f"{e.report['incomplete_mapping_count']}/{e.report['total_mappings']} "
                    "mapping(s) incomplete. Re-run with --completeness-report to persist details."
                )
            return 1
        _log_success(f"Extracted {len(lineage)} lineage records")

        if args.folder:
            before_count = len(lineage)
            lineage = [x for x in lineage if x["folder"] == args.folder]
            _log_info(f"Filtered by folder '{args.folder}': {before_count} -> {len(lineage)} records")

        by_mapping: dict[str, list[dict[str, Any]]] = defaultdict(list)
        for rec in lineage:
            by_mapping[rec["mapping"]].append(rec)
        _log_info(f"Grouped into {len(by_mapping)} unique mappings")

        out_dir = Path(args.lineage_out_dir or ".")
        out_dir.mkdir(parents=True, exist_ok=True)
        _log_success(f"Output directory ready: {out_dir.absolute()}")
        _log_info("-" * 70)

        written = 0
        mapping_names = _mapping_names_in_order(root)
        for idx_num, mapping_name in enumerate(mapping_names, 1):
            records = by_mapping.get(mapping_name, [])
            if not records:
                _log_warning(f"[{idx_num}/{len(mapping_names)}] {mapping_name}: No records found")
                continue

            _log_progress(idx_num, len(mapping_names), f"Writing {mapping_name}: {len(records)} records")

            serialisable = [
                {
                    **rec,
                    "source_nodes": [list(x) for x in rec["source_nodes"]],
                    "path_nodes": [list(x) for x in rec["path_nodes"]],
                }
                for rec in records
            ]

            out_file = out_dir / f"{mapping_name}_lineage_one_mapping.json"
            try:
                out_file.write_text(
                    json.dumps(serialisable, indent=2, ensure_ascii=False, default=str),
                    encoding="utf-8",
                )
                _log_success(f"  Wrote {out_file.name} ({len(records)} records)")
                written += 1
            except Exception as e:
                _log_error(f"  Failed to write {out_file.name}: {e}")

        _log_info("-" * 70)
        _log_success(f"All-mappings processing complete")
        _log_info(f"  Total records : {len(lineage)}")
        _log_info(f"  Files written : {written}/{len(by_mapping)}")

        if args.lineage_only:
            elapsed = time.time() - start_time
            _log_info("-" * 70)
            _log_success(f"Completed in {elapsed:.2f} seconds")
            return 0

    _log_info("=" * 70)
    _log_info("STANDARD MODE: Extracting lineage")
    _log_info("=" * 70)

    final_targets = args.final_targets if not args.all_targets else None
    auto_all_targets_for_mapping = bool(args.mapping and not args.all_targets and not args.final_targets)
    try:
        lineage = extract_target_bound_fields(
            data,
            final_target_definitions=final_targets,
            include_all_targets=(args.all_targets or auto_all_targets_for_mapping),
            checkpoint_dir=checkpoint_dir,
            enable_checkpoints=enable_checkpoints,
            enable_cache=enable_cache,
            strict_completeness=args.strict_completeness,
            completeness_report_path=args.completeness_report,
        )
    except IncompleteLineageError as e:
        _log_error(str(e))
        if args.completeness_report:
            _log_error(f"See completeness report for details: {args.completeness_report}")
        else:
            _log_error(
                f"{e.report['incomplete_mapping_count']}/{e.report['total_mappings']} "
                "mapping(s) incomplete. Re-run with --completeness-report to persist details."
            )
        return 1
    _log_success(f"Extracted {len(lineage)} initial lineage records")

    if args.folder:
        before_count = len(lineage)
        lineage = [x for x in lineage if x["folder"] == args.folder]
        _log_info(f"Filtered by folder '{args.folder}': {before_count} -> {len(lineage)} records")

    if args.mapping:
        before_count = len(lineage)
        lineage = [x for x in lineage if x["mapping"] == args.mapping]
        _log_info(f"Filtered by mapping '{args.mapping}': {before_count} -> {len(lineage)} records")

    _log_info("-" * 70)
    _log_success(f"Total records for output: {len(lineage)}")

    serialisable = [
        {
            **rec,
            "source_nodes": [list(x) for x in rec["source_nodes"]],
            "path_nodes": [list(x) for x in rec["path_nodes"]],
        }
        for rec in lineage
    ]

    # Write diagnostics if requested
    if args.diagnostics_out:
        # Re-run to collect diagnostics (or collect from last run if we refactor)
        _log_info(f"Writing diagnostics to {args.diagnostics_out}...")
        diag_entries = [
            rec.get("unresolved_nodes", []) for rec in lineage
        ]
        flat_diag = [item for sub in diag_entries for item in sub]
        try:
            Path(args.diagnostics_out).write_text(
                json.dumps(flat_diag, indent=2, ensure_ascii=False, default=str),
                encoding="utf-8",
            )
            _log_success(f"Diagnostics written: {args.diagnostics_out}")
        except Exception as e:
            _log_error(f"Failed to write diagnostics: {e}")

    if args.lineage_out:
        out_lineage = Path(args.lineage_out)
        out_lineage.parent.mkdir(parents=True, exist_ok=True)
        try:
            out_lineage.write_text(
                json.dumps(serialisable, indent=2, ensure_ascii=False, default=str),
                encoding="utf-8",
            )
            _log_success(f"Lineage JSON written: {out_lineage.absolute()}")
        except Exception as e:
            _log_error(f"Failed to write lineage JSON: {e}")
            return 1

    if args.lineage_only:
        print(json.dumps(serialisable, indent=2, ensure_ascii=False, default=str))
        elapsed = time.time() - start_time
        _log_info("-" * 70)
        _log_success(f"Completed in {elapsed:.2f} seconds")
        return 0

    _log_info("=" * 70)
    _log_info("BUILDING DOCX BRIEFING...")
    _log_info("=" * 70)

    try:
        out = build_briefing_docx(
            data,
            out_path=args.output,
            folder_name=args.folder,
            mapping_name=args.mapping,
            final_target_definitions=final_targets,
            include_all_targets=(args.all_targets or auto_all_targets_for_mapping),
            checkpoint_dir=checkpoint_dir,
            enable_checkpoints=enable_checkpoints,
            enable_cache=enable_cache,
            strict_completeness=args.strict_completeness,
            completeness_report_path=args.completeness_report,
        )
        _log_success(f"DOCX Briefing created: {out.absolute()} ({out.stat().st_size:,} bytes)")
    except IncompleteLineageError as e:
        _log_error(str(e))
        if args.completeness_report:
            _log_error(f"See completeness report for details: {args.completeness_report}")
        else:
            _log_error(
                f"{e.report['incomplete_mapping_count']}/{e.report['total_mappings']} "
                "mapping(s) incomplete. Re-run with --completeness-report to persist details."
            )
        return 1
    except Exception as e:
        _log_error(f"Failed to build DOCX briefing: {e}")
        return 1

    elapsed = time.time() - start_time
    _log_info("#" * 70)
    _log_success(f"PROCESS COMPLETE — Elapsed: {elapsed:.2f}s")
    _log_info("#" * 70)
    return 0


if __name__ == "__main__":
    sys.exit(_cli())