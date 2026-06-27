import json
import os
import subprocess
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ============================================================
# INPUT / OUTPUT
# ============================================================

if len(sys.argv) == 3:
    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])

    if not input_path.exists():
        raise ValueError(f"Input path not found: {input_path}")

    # Single-file mode: python 05_overview_to_docs.py input.json output_dir
    # Optional: output_path can be an explicit .docx file path.
    if input_path.is_file():
        if input_path.suffix.lower() != ".json":
            raise ValueError(f"Input file must be a JSON file: {input_path}")

        if output_path.suffix.lower() == ".docx":
            out_docx = output_path
            out_docx.parent.mkdir(parents=True, exist_ok=True)
        else:
            output_path.mkdir(parents=True, exist_ok=True)
            out_docx = output_path / f"{input_path.stem}.docx"

        env = os.environ.copy()
        env["OVERVIEW_TO_DOCS_SRC"] = str(input_path)
        env["OVERVIEW_TO_DOCS_OUT"] = str(out_docx)
        subprocess.run(
            [sys.executable, __file__],
            check=True,
            env=env,
        )

        print(json.dumps({
            "mode": "single",
            "input_file": str(input_path),
            "output_file": str(out_docx),
        }, indent=2))
        sys.exit(0)

    if not input_path.is_dir():
        raise ValueError(f"Input path is neither file nor directory: {input_path}")

    output_path.mkdir(parents=True, exist_ok=True)

    json_files = sorted(input_path.glob("*.json"))
    if not json_files:
        raise ValueError(f"No JSON files found in input directory: {input_path}")

    converted = 0
    for json_file in json_files:
        out_docx = output_path / f"{json_file.stem}.docx"
        env = os.environ.copy()
        env["OVERVIEW_TO_DOCS_SRC"] = str(json_file)
        env["OVERVIEW_TO_DOCS_OUT"] = str(out_docx)

        subprocess.run(
            [sys.executable, __file__],
            check=True,
            env=env,
        )
        converted += 1

    print(json.dumps({
        "mode": "batch",
        "input_directory": str(input_path),
        "output_directory": str(output_path),
        "json_files_converted": converted,
    }, indent=2))
    sys.exit(0)

if len(sys.argv) not in (1,):
    raise ValueError(
        "Usage: python 05_overview_to_docs.py <input_directory_path> <output_directory>"
    )

SRC = os.environ.get(
    "OVERVIEW_TO_DOCS_SRC",
    "input_json/WF_DTM_JPRFXX_OTC_VAL_runtime_output.json",
)
OUT = os.environ.get(
    "OVERVIEW_TO_DOCS_OUT",
    "docs/JSFA_OTC_VAL/WF_DTM_JPRFXX_OTC_VAL_Overview_Workflow.docx",
)


# ============================================================
# LOAD JSON
# ============================================================

with open(SRC, "r", encoding="utf-8") as file:
    data = json.load(file)

if not isinstance(data, dict):
    raise ValueError("Expected JSON root to be a dictionary/object.")


# ============================================================
# HELPER FUNCTIONS
# ============================================================

def clean_text(value, max_chars=None):
    """
    Cleans text for Word table cells.
    Preserves meaningful line breaks and trims unnecessary whitespace.
    """
    if value is None:
        text = ""
    elif isinstance(value, (dict, list)):
        text = json.dumps(value, ensure_ascii=False)
    else:
        text = str(value)

    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = "\n".join(
        line.strip()
        for line in text.split("\n")
        if line.strip()
    )

    if max_chars and len(text) > max_chars:
        text = text[:max_chars - 24] + " ... [truncated in cell]"

    return text


def shade(cell, fill="D9EAF7"):
    """
    Applies background color to a table cell.
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=7):
    """
    Sets formatted text inside a Word table cell.
    """
    cell.text = ""
    paragraph = cell.paragraphs[0]

    parts = clean_text(text).split("\n") if text is not None else [""]

    if not parts:
        parts = [""]

    for index, part in enumerate(parts):
        if index:
            paragraph.add_run().add_break()

        run = paragraph.add_run(part)
        run.bold = bold
        run.font.name = "Aptos"
        run.font.size = Pt(size)

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def repeat_header(row):
    """
    Repeats table header row across pages in Word.
    """
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc, headers, rows, font_size=7):
    """
    Adds a formatted table to the Word document.
    """
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    repeat_header(table.rows[0])

    # Header row
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, size=7)
        shade(table.rows[0].cells[index])

    # Data rows
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value, size=font_size)

    return table


def get_table_rows(records, headers):
    """
    Converts list of dictionaries into ordered rows based on headers.
    """
    rows = []

    for record in records:
        row = []
        for header in headers:
            row.append(record.get(header, ""))
        rows.append(row)

    return rows


# ============================================================
# EXTRACT JSON SECTIONS
# ============================================================

workflow_session_mapping_transformation_table = data.get(
    "workflow_session_mapping_transformation_table",
    []
)

mapping_flowcharts = data.get(
    "mapping_flowcharts",
    []
)

section_3_session_source_target_table = data.get(
    "section_3_session_source_target_table",
    []
)

section_3_session_execution_flowchart = data.get(
    "section_3_session_execution_flowchart",
    ""
)

section_3_session_table_flowchart = data.get(
    "section_3_session_table_flowchart",
    ""
)

section_4_writer_export_table = data.get(
    "section_4_writer_export_table",
    []
)


# ============================================================
# CREATE WORD DOCUMENT - VERSION 2 COMPACT LANDSCAPE FORMAT
# ============================================================

doc = Document()

section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Inches(11.69)
section.page_height = Inches(8.27)
section.top_margin = Inches(0.45)
section.bottom_margin = Inches(0.45)
section.left_margin = Inches(0.45)
section.right_margin = Inches(0.45)

for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
    doc.styles[style_name].font.name = "Aptos"

doc.styles["Normal"].font.size = Pt(8)


# ============================================================
# SECTION 1 - WORKFLOW SESSION MAPPING TRANSFORMATION TABLE
# ============================================================

doc.add_heading(
    "1. Workflow Session Mapping Transformation Table",
    level=1
)

workflow_headers = [
    "Workflow Name",
    "Session Name",
    "Mapping Name",
    "Transformation Name"
]

workflow_rows = get_table_rows(
    workflow_session_mapping_transformation_table,
    workflow_headers
)

if workflow_rows:
    add_table(
        doc,
        workflow_headers,
        workflow_rows,
        font_size=7
    )
else:
    doc.add_paragraph(
        "workflow_session_mapping_transformation_table: "
        "Not explicitly defined in JSON."
    )


# ============================================================
# SECTION 2 - MAPPING FLOWCHARTS
# ============================================================

doc.add_heading(
    "2. Mapping Flowcharts",
    level=1
)

flowchart_rows = []

for item in mapping_flowcharts:
    mapping_name = item.get("Mapping Heading", "")
    flowchart_text = item.get("Flowchart", "")

    flowchart_rows.append([
        mapping_name,
        flowchart_text
    ])

if flowchart_rows:
    add_table(
        doc,
        [
            "Mapping Name",
            "Flowchart"
        ],
        flowchart_rows,
        font_size=7
    )
else:
    doc.add_paragraph(
        "mapping_flowcharts: Not explicitly defined in JSON."
    )


# ============================================================
# SECTION 3 - SESSION SOURCE TARGET TABLE & FLOWCHARTS
# ============================================================

doc.add_heading(
    "3. Session Source & Target Tables",
    level=1
)

if section_3_session_source_target_table:
    session_headers = [
        "Session Name",
        "Mapping Name",
        "Source Table",
        "Target Table"
    ]

    session_rows = get_table_rows(
        section_3_session_source_target_table,
        session_headers
    )

    add_table(
        doc,
        session_headers,
        session_rows,
        font_size=7
    )
else:
    doc.add_paragraph(
        "section_3_session_source_target_table: Not defined in JSON."
    )

# Session Execution Flowchart
doc.add_heading(
    "3.1 Session Execution Flowchart (Runtime Order)",
    level=2
)

if section_3_session_execution_flowchart:
    p = doc.add_paragraph()
    run = p.add_run(clean_text(section_3_session_execution_flowchart))
    run.font.name = "Aptos"
    run.font.size = Pt(8)
else:
    doc.add_paragraph(
        "session_execution_flowchart: Not defined in JSON."
    )

# Session Table Flowchart
doc.add_heading(
    "3.2 Session Table Flowchart (Data Flow)",
    level=2
)

if section_3_session_table_flowchart:
    p = doc.add_paragraph()
    run = p.add_run(clean_text(section_3_session_table_flowchart))
    run.font.name = "Aptos"
    run.font.size = Pt(8)
else:
    doc.add_paragraph(
        "session_table_flowchart: Not defined in JSON."
    )


# ============================================================
# SECTION 4 - WRITER EXPORT TABLE (IF AVAILABLE)
# ============================================================

if section_4_writer_export_table:
    doc.add_heading(
        "4. File Writer Exports (Reporting)",
        level=1
    )

    export_headers = [
        "Target Table",
        "Export Filename",
        "Export Directory",
        "Export Type",
        "Purpose"
    ]

    export_rows = get_table_rows(
        section_4_writer_export_table,
        export_headers
    )

    if export_rows:
        add_table(
            doc,
            export_headers,
            export_rows,
            font_size=7
        )
    else:
        doc.add_paragraph(
            "section_4_writer_export_table: No writer exports found."
        )


# ============================================================
# SAVE DOCUMENT
# ============================================================

# Remove existing file if it exists to avoid permission errors
if os.path.exists(OUT):
    try:
        os.remove(OUT)
    except PermissionError:
        # If we can't remove it, try to add a timestamp to the filename
        base_name = Path(OUT).stem
        ext = Path(OUT).suffix
        parent_dir = Path(OUT).parent
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        OUT = str(parent_dir / f"{base_name}_{timestamp}{ext}")

doc.save(OUT)


# ============================================================
# VALIDATION
# ============================================================

check_doc = Document(OUT)
check_section = check_doc.sections[0]

stats = {
    "output_file": OUT,
    "landscape_verified": check_section.page_width > check_section.page_height,
    "page_width_inches": round(check_section.page_width / 914400, 2),
    "page_height_inches": round(check_section.page_height / 914400, 2),
    "workflow_records": len(workflow_session_mapping_transformation_table),
    "mapping_flowchart_records": len(mapping_flowcharts),
    "session_source_target_records": len(section_3_session_source_target_table),
    "writer_export_records": len(section_4_writer_export_table),
    "tables_created": len(check_doc.tables),
    "paragraphs_created": len(check_doc.paragraphs),
    "file_size_bytes": os.path.getsize(OUT),
    "format_followed": "Four-section business format (Sections 1-4)"
}

print(json.dumps(stats, indent=2))