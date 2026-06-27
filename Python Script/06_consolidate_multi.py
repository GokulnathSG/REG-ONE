"""
06_consolidate_multi.py
=======================
Consolidate multiple JSON files into one JSON, or multiple DOCX files into one DOCX,
without changing existing project scripts.

Examples
--------
# Consolidate all JSON files in a folder
python 06_consolidate_multi.py --mode json \
    -i output/JSON/PUMA_01 \
    -o output/PUMA_01_consolidated.json

# Consolidate specific JSON files in order
python 06_consolidate_multi.py --mode json \
    -i file1.json file2.json file3.json \
    -o output/consolidated.json

# Consolidate all DOCX files in a folder
python 06_consolidate_multi.py --mode doc \
    -i docs/Puma_01 \
    -o docs/Puma_01_master.docx

# Consolidate specific DOCX files in order
python 06_consolidate_multi.py --mode doc \
    -i docs/Puma_01/a.docx docs/Puma_01/b.docx \
    -o docs/Master_FSD.docx
"""

from __future__ import annotations

import argparse
import json
import sys
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.shared import Pt

try:
    from docxcompose.composer import Composer
except ImportError:
    Composer = None


def _resolve_input_files(inputs: list[str], extension: str) -> list[Path]:
    """Resolve files from folders/files and filter by extension."""
    resolved: list[Path] = []

    for raw in inputs:
        p = Path(raw)
        if p.is_dir():
            resolved.extend(sorted(p.glob(f"*{extension}")))
        elif p.is_file() and p.suffix.lower() == extension:
            resolved.append(p)
        else:
            print(f"[WARN] Skipping invalid input: {raw}", file=sys.stderr)

    # Remove duplicates while preserving first-seen order
    deduped: list[Path] = []
    seen: set[str] = set()
    for p in resolved:
        key = str(p.resolve())
        if key not in seen:
            seen.add(key)
            deduped.append(p)

    return deduped


def _ensure_parent_dir(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)


def _set_landscape_all_sections(doc: Document) -> None:
    """Force landscape orientation for all sections in a document."""
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Pt(841.89)   # A4 width in landscape
        section.page_height = Pt(595.28)  # A4 height in landscape


def consolidate_json(input_files: list[Path], output_file: Path) -> None:
    """Create a single JSON output containing all input JSON payloads."""
    bundled_inputs: list[dict] = []
    merged_records: list = []
    all_roots_are_lists = True

    for file_path in input_files:
        with file_path.open("r", encoding="utf-8") as f:
            data = json.load(f)

        bundled_inputs.append(
            {
                "file": file_path.name,
                "path": str(file_path),
                "root_type": type(data).__name__,
                "data": data,
            }
        )

        if isinstance(data, list):
            merged_records.extend(data)
        else:
            all_roots_are_lists = False

    out_obj = {
        "generated_at_utc": datetime.now(timezone.utc).isoformat(),
        "input_count": len(input_files),
        "inputs": bundled_inputs,
    }

    # Convenience field for common lineage files where each input is a list.
    if all_roots_are_lists:
        out_obj["merged_records"] = merged_records
        out_obj["merged_record_count"] = len(merged_records)

    _ensure_parent_dir(output_file)
    with output_file.open("w", encoding="utf-8") as f:
        json.dump(out_obj, f, indent=2, ensure_ascii=False)

    print(f"[OK] JSON consolidated: {output_file}")
    print(f"[OK] Input files: {len(input_files)}")
    if all_roots_are_lists:
        print(f"[OK] Merged records: {len(merged_records)}")


def _append_doc_body(dst: Document, src: Document) -> None:
    """Append body XML from src doc into dst doc (except section properties)."""
    for element in src.element.body:
        if element.tag.endswith("sectPr"):
            continue
        dst.element.body.append(deepcopy(element))


def _build_separator_doc(index: int, file_name: str) -> Document:
    sep_doc = Document()
    p = sep_doc.add_paragraph(f"Source {index}: {file_name}")
    p.runs[0].bold = True
    _set_landscape_all_sections(sep_doc)
    return sep_doc


def consolidate_docx(input_files: list[Path], output_file: Path, title: str | None) -> None:
    """Create a single DOCX by appending all source DOCX documents in order."""
    _ensure_parent_dir(output_file)

    if Composer is not None:
        # Composer preserves source document details (tables, images, numbering, references)
        # more reliably than manual XML append.
        master = Document()
        _set_landscape_all_sections(master)

        if title:
            p = master.add_paragraph(title)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(16)
            master.add_paragraph()

        composer = Composer(master)
        for index, file_path in enumerate(input_files, start=1):
            composer.append(_build_separator_doc(index, file_path.name))
            composer.append(Document(str(file_path)))
        composer.save(str(output_file))
    else:
        # Fallback path if docxcompose is not installed.
        print("[WARN] docxcompose not installed; using fallback append mode.", file=sys.stderr)
        master = Document()

        if title:
            p = master.add_paragraph(title)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(16)
            master.add_paragraph()

        for index, file_path in enumerate(input_files, start=1):
            src = Document(str(file_path))
            sep = master.add_paragraph(f"Source {index}: {file_path.name}")
            sep.runs[0].bold = True
            _append_doc_body(master, src)
            if index != len(input_files):
                master.add_page_break()

        master.save(str(output_file))

    # Enforce landscape orientation in final output document.
    out_doc = Document(str(output_file))
    _set_landscape_all_sections(out_doc)
    out_doc.save(str(output_file))

    print(f"[OK] DOCX consolidated: {output_file}")
    print(f"[OK] Input files: {len(input_files)}")


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Consolidate multiple JSON or DOCX files into a single output file.",
    )
    parser.add_argument(
        "--mode",
        required=True,
        choices=["json", "doc"],
        help="Consolidation mode: json or doc",
    )
    parser.add_argument(
        "-i",
        "--input",
        nargs="+",
        required=True,
        metavar="PATH",
        help="Input files and/or folders",
    )
    parser.add_argument(
        "-o",
        "--output",
        required=True,
        metavar="PATH",
        help="Output file path",
    )
    parser.add_argument(
        "--title",
        default=None,
        help="Optional title for consolidated DOCX output",
    )
    return parser


def main() -> None:
    args = build_parser().parse_args()

    mode = args.mode
    extension = ".json" if mode == "json" else ".docx"

    input_files = _resolve_input_files(args.input, extension)
    if not input_files:
        print(f"[ERROR] No valid {extension} inputs found.", file=sys.stderr)
        sys.exit(1)

    output_file = Path(args.output)

    print(f"[INFO] Mode: {mode}")
    print(f"[INFO] Inputs detected: {len(input_files)}")
    for f in input_files:
        print(f"       - {f}")

    if mode == "json":
        consolidate_json(input_files, output_file)
    else:
        consolidate_docx(input_files, output_file, args.title)


if __name__ == "__main__":
    main()
