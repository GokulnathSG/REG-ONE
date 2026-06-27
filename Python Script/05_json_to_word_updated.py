import json
import os
import subprocess
import sys
from collections import defaultdict, Counter
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ============================================================
# INPUT / OUTPUT FILES
# ============================================================

if len(sys.argv) == 3:
    input_dir = Path(sys.argv[1])
    output_dir = Path(sys.argv[2])

    if not input_dir.exists() or not input_dir.is_dir():
        raise ValueError(f"Input directory not found: {input_dir}")

    output_dir.mkdir(parents=True, exist_ok=True)

    json_files = sorted(input_dir.glob("*.json"))
    if not json_files:
        raise ValueError(f"No JSON files found in input directory: {input_dir}")

    converted = 0
    for json_file in json_files:
        out_docx = output_dir / f"{json_file.stem}.docx"
        env = os.environ.copy()
        env["JSON_TO_WORD_SRC"] = str(json_file)
        env["JSON_TO_WORD_OUT"] = str(out_docx)

        subprocess.run(
            [sys.executable, __file__],
            check=True,
            env=env,
        )
        converted += 1

    print(json.dumps({
        "mode": "batch",
        "input_directory": str(input_dir),
        "output_directory": str(output_dir),
        "json_files_converted": converted,
    }, indent=2))
    sys.exit(0)

if len(sys.argv) not in (1,):
    raise ValueError(
        "Usage: python 03_json_to_word.py <input_directory_path> <output_directory>"
    )

SRC = os.environ.get("JSON_TO_WORD_SRC", "runtime_output.json")
OUT = os.environ.get("JSON_TO_WORD_OUT", "runtime_output.docx")


# ============================================================
# LOAD JSON
# ============================================================

with open(SRC, "r", encoding="utf-8") as f:
    data = json.load(f)

if not isinstance(data, list):
    raise ValueError("Expected JSON root to be a list of lineage records")


# ============================================================
# HELPER FUNCTIONS
# ============================================================

def clean_text(value, max_chars=None):
    """
    Cleans text for Word table cells.
    Keeps line breaks compact.
    Optionally truncates long values to avoid oversized Word cells.
    """
    if value is None:
        return ""

    if isinstance(value, (list, dict)):
        text = json.dumps(value, ensure_ascii=False)
    else:
        text = str(value)

    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = "\n".join(
        line.strip()
        for line in text.split("\n")
        if line.strip()
    )

    return text


def join_nodes(nodes, max_chars=None):
    """
    Converts path_nodes into readable lineage path:
    Source.Field → Transformation.Field → Target.Field
    """
    if not nodes:
        return "Not explicitly traceable in JSON"

    output = []

    for node in nodes:
        if isinstance(node, list) and len(node) >= 2:
            output.append(f"{node[0]}.{node[1]}")
        else:
            output.append(clean_text(node))

    return clean_text(" → ".join(output), max_chars=max_chars)


def join_list(items, empty="Not explicitly defined in JSON", max_chars=None):
    """
    Converts a list into newline-separated text.
    """
    if not items:
        return empty

    return clean_text(
        "\n".join(clean_text(item) for item in items),
        max_chars=max_chars
    )


def join_path_list(items, empty="Not explicitly traceable in JSON", max_chars=None):
    """
    Converts a list of node strings into a readable lineage path.
    """
    if not items:
        return empty

    return clean_text(" -> ".join(clean_text(item) for item in items), max_chars=max_chars)


def shade(cell, fill="D9EAF7"):
    """
    Applies background shading to table header cells.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_text(cell, text, bold=False, size=7):
    """
    Sets formatted text inside a Word table cell.
    """
    cell.text = ""
    paragraph = cell.paragraphs[0]

    parts = clean_text(text).split("\n") if text is not None else [""]

    if not parts:
        parts = [""]

    for index, part in enumerate(parts):
        if index:
            paragraph.add_run().add_break()

        run = paragraph.add_run(part)
        run.bold = bold
        run.font.name = "Aptos"
        run.font.size = Pt(size)

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def repeat_header(row):
    """
    Repeats table header row across pages.
    """
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def add_table(doc, headers, rows, font_size=7):
    """
    Adds a styled table to the Word document.
    """
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    repeat_header(table.rows[0])

    # Header row
    for i, header in enumerate(headers):
        set_text(table.rows[0].cells[i], header, bold=True, size=7)
        shade(table.rows[0].cells[i])

    # Data rows
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_text(cells[i], value, size=font_size)

    return table


# ============================================================
# METADATA EXTRACTION
# ============================================================

folders = sorted({
    record.get("folder", "")
    for record in data
    if isinstance(record, dict) and record.get("folder")
})

mappings = sorted({
    record.get("mapping", "")
    for record in data
    if isinstance(record, dict) and record.get("mapping")
})

target_instances = sorted({
    record.get("target_instance", "")
    for record in data
    if isinstance(record, dict) and record.get("target_instance")
})

transform_types = Counter()
lookup_tables = set()

for record in data:
    if not isinstance(record, dict):
        continue

    for transformation in record.get("transformations_touched") or []:
        transform_types[
            transformation.get("transform_type") or "Not specified"
        ] += 1

    for lookup in record.get("lookup_used") or []:
        if lookup.get("lookup_table"):
            lookup_tables.add(lookup.get("lookup_table"))


# ============================================================
# CREATE WORD DOCUMENT
# ============================================================

doc = Document()

section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Inches(11.69)
section.page_height = Inches(8.27)
section.top_margin = Inches(0.45)
section.bottom_margin = Inches(0.45)
section.left_margin = Inches(0.45)
section.right_margin = Inches(0.45)

for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
    doc.styles[style_name].font.name = "Aptos"

doc.styles["Normal"].font.size = Pt(8)


# ============================================================
# TITLE
# ============================================================

paragraph = doc.add_paragraph()
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

run = paragraph.add_run("Target Field Lineage Documentation")
run.bold = True
run.font.size = Pt(18)

paragraph = doc.add_paragraph()
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

run = paragraph.add_run(
    "Converted from JSON to Word document — Landscape Orientation"
)
run.italic = True
run.font.size = Pt(10)


# ============================================================
# SECTION 1 - JSON CONVERSION SUMMARY
# ============================================================

doc.add_heading("1. JSON Conversion Summary", level=1)

summary_rows = [
    ["Source File", SRC],
    ["Orientation", "Landscape"],
    [
        "Folder(s)",
        ", ".join(folders) if folders else "Not explicitly defined in JSON"
    ],
    ["Total Lineage Records", str(len(data))],
    ["Mappings", f"{len(mappings)}: " + ", ".join(mappings)],
    [
        "Target Instances",
        f"{len(target_instances)}: " + ", ".join(target_instances)
    ],
    [
        "Transformation Types",
        ", ".join(
            f"{key} ({value})"
            for key, value in transform_types.most_common()
        ) or "Not explicitly defined in JSON"
    ],
    [
        "Lookup Tables",
        ", ".join(sorted(lookup_tables))
        if lookup_tables
        else "Not explicitly defined in JSON"
    ],
    ["Generated On", datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
]

add_table(
    doc,
    ["Item", "Value"],
    summary_rows,
    font_size=8
)


# ============================================================
# SECTION 2 - MAPPING AND TARGET SUMMARY
# ============================================================

doc.add_heading("2. Mapping and Target Summary", level=1)

mapping_rows = []

for mapping in mappings:
    records = [
        record
        for record in data
        if isinstance(record, dict) and record.get("mapping") == mapping
    ]

    mapping_rows.append([
        mapping,
        ", ".join(sorted({
            record.get("target_instance", "")
            for record in records
            if record.get("target_instance")
        })),
        str(len(records)),
        str(sum(
            1 for record in records
            if clean_text(record.get("transformation_rule"))
        )),
        str(sum(
            1 for record in records
            if record.get("lookup_used")
        )),
    ])

add_table(
    doc,
    [
        "Mapping",
        "Target Instance(s)",
        "Target Field Count",
        "Fields with Transformation Rule",
        "Fields using Lookup"
    ],
    mapping_rows,
    font_size=7
)


# ============================================================
# SECTION 3 - DETAILED TARGET FIELD LINEAGE
# ============================================================

doc.add_heading("3. Detailed Target Field Lineage", level=1)

grouped_records = defaultdict(list)

for record in data:
    if isinstance(record, dict):
        grouped_records[
            (
                record.get("mapping", ""),
                record.get("target_instance", "")
            )
        ].append(record)


for (mapping, target_instance), records in sorted(
    grouped_records.items(),
    key=lambda item: (item[0][0], item[0][1])
):
    doc.add_heading(f"Mapping: {mapping}", level=2)
    doc.add_heading(f"Target Instance: {target_instance}", level=3)

    lineage_rows = []

    for index, record in enumerate(records, start=1):

        transformations = []

        for transformation in record.get("transformations_touched") or []:
            item = (
                f"{transformation.get('instance_name', '')} "
                f"({transformation.get('transform_type', '')})"
            )

            if transformation.get("field_name"):
                item += f".{transformation.get('field_name')}"

            if transformation.get("expression"):
                item += f" = {transformation.get('expression')}"

            transformations.append(item)

        lookups = []

        for lookup in record.get("lookup_used") or []:
            lookup_parts = []

            if lookup.get("transform_name"):
                lookup_parts.append(
                    f"Transform: {lookup.get('transform_name')}"
                )

            if lookup.get("lookup_table"):
                lookup_parts.append(
                    f"Table: {lookup.get('lookup_table')}"
                )

            if lookup.get("lookup_condition"):
                lookup_parts.append(
                    f"Condition: {lookup.get('lookup_condition')}"
                )

            if lookup.get("lookup_sql_override"):
                lookup_parts.append(
                    f"SQL Override: {lookup.get('lookup_sql_override')}"
                )

            lookups.append("\n".join(lookup_parts))

        lineage_rows.append([
            str(index),
            clean_text(
                record.get("target_field")
                or record.get("target_column_name")
            ),
            join_list(
                record.get("ultimate_source_fields") or [],
                "Not explicitly traceable in JSON"
            ),
            join_nodes(
                record.get("path_nodes") or []
            ),
            join_list(
                transformations,
                "No transformation touched"
            ),
            clean_text(
                record.get("transformation_rule")
            )
            or clean_text(
                record.get("transformation_rule_plain_english")
            )
            or "Direct pass-through",
            join_list(
                lookups,
                "Not used"
            ),
            join_list(
                record.get("filter_or_router_conditions") or [],
                "Not defined"
            ),
        ])

    add_table(
        doc,
        [
            "#",
            "Target Field",
            "Ultimate Source Field(s)",
            "Lineage Path",
            "Transformation(s) Touched",
            "Rule / Derivation",
            "Lookup Used",
            "Filter / Router Conditions"
        ],
        lineage_rows,
        font_size=6
    )


# ============================================================
# SECTION 3A - EXPRESSION ATTRIBUTE RECURSIVE LINEAGE
# ============================================================

doc.add_heading("3A. Expression Attribute Recursive Lineage", level=1)

expr_grouped_rows = defaultdict(list)
expr_seen_keys = set()

for record in data:
    if not isinstance(record, dict):
        continue

    mapping = record.get("mapping", "")
    target_instance = record.get("target_instance", "")
    parent_target_field = clean_text(
        record.get("target_field") or record.get("target_column_name")
    )
    parent_expression = clean_text(record.get("transformation_rule"))

    for expr_block in record.get("expression_field_lineage_section") or []:
        expression_used = clean_text(expr_block.get("expression"))
        expression_owner = clean_text(
            f"{expr_block.get('expression_instance', '')}.{expr_block.get('expression_field', '')}".strip(".")
        )

        for field_lineage in expr_block.get("field_references_lineage") or []:
            field_instance = clean_text(field_lineage.get("field_instance"))
            field_name = clean_text(field_lineage.get("field_name"))
            target_field_ref = ".".join(part for part in [field_instance, field_name] if part)

            dedupe_key = (
                mapping,
                target_instance,
                parent_target_field,
                expression_owner,
                expression_used,
                target_field_ref,
                tuple(field_lineage.get("lineage_path") or []),
            )
            if dedupe_key in expr_seen_keys:
                continue
            expr_seen_keys.add(dedupe_key)

            rule_text = clean_text(field_lineage.get("rule")) or "Direct pass-through"
            if expression_used:
                rule_text = (
                    f"Used in expression ({expression_owner or 'unknown'}): {expression_used}\n"
                    f"Derived rule: {rule_text}"
                )

            expr_grouped_rows[(mapping, target_instance)].append([
                target_field_ref or "Not explicitly defined in JSON",
                join_list(
                    field_lineage.get("ultimate_source_fields") or [],
                    "Not explicitly traceable in JSON",
                ),
                join_path_list(
                    field_lineage.get("lineage_path") or [],
                    "Not explicitly traceable in JSON",
                ),
                join_list(
                    field_lineage.get("transformations_touched") or [],
                    "No transformation touched",
                ),
                clean_text(rule_text),
                join_list(
                    record.get("lookup_used") or [],
                    "Not used",
                ),
                join_list(
                    record.get("filter_or_router_conditions") or [],
                    "Not defined",
                ),
            ])

if expr_grouped_rows:
    for (mapping, target_instance), expr_rows in sorted(
        expr_grouped_rows.items(),
        key=lambda item: (item[0][0], item[0][1]),
    ):
        doc.add_heading(f"Mapping: {mapping}", level=2)
        doc.add_heading(f"Target Instance: {target_instance}", level=3)

        add_table(
            doc,
            [
                "Target Field",
                "Ultimate Source Field(s)",
                "Lineage Path",
                "Transformation(s) Touched",
                "Rule / Derivation",
                "Lookup Used",
                "Filter / Router Conditions",
            ],
            expr_rows,
            font_size=6,
        )
else:
    doc.add_paragraph(
        "No expression-referenced attribute lineage found in JSON."
    )


# ============================================================
# SECTION 4 - EXPLICIT TRANSFORMATION RULES
# ============================================================

doc.add_heading("4. Appendix - Explicit Transformation Rules", level=1)

rule_rows = []

for record in data:
    if not isinstance(record, dict):
        continue

    rule = clean_text(record.get("transformation_rule"))

    if rule:
        rule_rows.append([
            record.get("mapping", ""),
            record.get("target_instance", ""),
            record.get("target_field", ""),
            clean_text(rule, max_chars=1800),
            clean_text(
                record.get("transformation_rule_plain_english"),
                max_chars=1200
            ),
        ])

if rule_rows:
    add_table(
        doc,
        [
            "Mapping",
            "Target Instance",
            "Target Field",
            "Transformation Rule",
            "Plain-English Rule"
        ],
        rule_rows,
        font_size=6
    )
else:
    doc.add_paragraph("No explicit transformation rules found in JSON.")


# ============================================================
# SECTION 5 - LOOKUP DETAILS
# ============================================================

doc.add_heading("5. Appendix - Lookup Details", level=1)

lookup_rows = []
seen_lookup_rows = set()

for record in data:
    if not isinstance(record, dict):
        continue

    for lookup in record.get("lookup_used") or []:
        key = (
            record.get("mapping", ""),
            record.get("target_field", ""),
            lookup.get("transform_name", ""),
            lookup.get("lookup_table", ""),
            lookup.get("lookup_condition", ""),
            lookup.get("lookup_sql_override", "")
        )

        if key in seen_lookup_rows:
            continue

        seen_lookup_rows.add(key)

        lookup_rows.append([
            record.get("mapping", ""),
            record.get("target_field", ""),
            lookup.get("transform_name", ""),
            lookup.get("lookup_table", ""),
            clean_text(
                lookup.get("lookup_condition"),
                max_chars=1000
            ),
            clean_text(
                lookup.get("lookup_sql_override"),
                max_chars=2200
            ),
        ])

if lookup_rows:
    add_table(
        doc,
        [
            "Mapping",
            "Target Field",
            "Lookup Transform",
            "Lookup Table",
            "Lookup Condition",
            "SQL Override"
        ],
        lookup_rows,
        font_size=6
    )
else:
    doc.add_paragraph("No lookup details found in JSON.")


# ============================================================
# SAVE DOCUMENT
# ============================================================

doc.save(OUT)


# ============================================================
# VALIDATION
# ============================================================

check_doc = Document(OUT)
check_section = check_doc.sections[0]

stats = {
    "output_file": OUT,
    "records_loaded": len(data),
    "mappings": len(mappings),
    "target_instances": len(target_instances),
    "tables_created": len(check_doc.tables),
    "paragraphs_created": len(check_doc.paragraphs),
    "page_width_inches": round(check_section.page_width / 914400, 2),
    "page_height_inches": round(check_section.page_height / 914400, 2),
    "landscape_verified": check_section.page_width > check_section.page_height,
    "raw_json_snapshot_included": False,
    "format_followed": "Version-2 compact landscape format"
}

print(json.dumps(stats, indent=2))